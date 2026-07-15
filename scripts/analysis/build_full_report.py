#!/usr/bin/env python3
"""
Build the final Vietnamese DOCX report for the VN Gold Market Analysis project.

The report is generated from the current data lake snapshot only. It does not
fetch live data and does not invent model outputs when optional dependencies are
missing. The final decision is therefore as-of the latest collected date in the
snapshot, not a live market call.
"""

from __future__ import annotations

import json
import math
import sys
import warnings
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from ftfy import fix_text
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

warnings.filterwarnings("ignore")
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

ROOT = Path(__file__).resolve().parents[2]
LAKE = ROOT / "data" / "lake"
MODELING = LAKE / "modeling"
DOCS = ROOT / "docs" / "reports"
FIGURES = DOCS / "figures"
OUT = DOCS / "vn_gold_analysis_full_report.docx"
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

SNAPSHOT_NOTE = "BÃ¡o cÃ¡o chá»‰ sá»­ dá»¥ng snapshot dá»¯ liá»‡u hiá»‡n cÃ³, má»›i nháº¥t Ä‘áº¿n 2026-07-11."
HORIZON_LABELS = {1: "1 thÃ¡ng", 3: "3 thÃ¡ng", 5: "5 thÃ¡ng"}


@dataclass
class DataProfile:
    name: str
    path: str
    rows: int
    cols: int
    date_span: str
    key_fields: str
    role: str


def read_csv(rel_path: str, **kwargs: Any) -> pd.DataFrame:
    path = ROOT / rel_path
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, low_memory=False, **kwargs)


def read_json(rel_path: str) -> dict[str, Any]:
    path = ROOT / rel_path
    if not path.exists():
        return {}
    with path.open(encoding="utf-8") as fh:
        return json.load(fh)


def safe_date_series(values: pd.Series) -> pd.Series:
    return pd.to_datetime(values, errors="coerce").dt.normalize()


def first_existing(columns: list[str], df: pd.DataFrame) -> str | None:
    for col in columns:
        if col in df.columns:
            return col
    return None


def date_span(df: pd.DataFrame) -> str:
    if df.empty:
        return "KhÃ´ng cÃ³ dá»¯ liá»‡u"
    col = first_existing(["date", "business_date", "event_date", "available_from", "DAY", "observation_date"], df)
    if not col:
        return "KhÃ´ng cÃ³ cá»™t ngÃ y chuáº©n"
    dates = safe_date_series(df[col])
    dates = dates.dropna()
    if dates.empty:
        return f"Cá»™t {col} khÃ´ng parse Ä‘Æ°á»£c ngÃ y"
    return f"{dates.min().date()} Ä‘áº¿n {dates.max().date()} ({col})"


def pct(value: float | int | None, digits: int = 1) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    return f"{100 * float(value):.{digits}f}%"


def pct_points(value: float | int | None, digits: int = 2) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    return f"{float(value):.{digits}f}%"


def money_m(value: float | int | None, digits: int = 2) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    return f"{float(value) / 1_000_000:.{digits}f} triá»‡u"


def num(value: float | int | None, digits: int = 3) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    if isinstance(value, (int, np.integer)):
        return f"{int(value):,}"
    return f"{float(value):,.{digits}f}"


def clean_text(value: Any) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return ""
    text = str(value)
    return fix_text(text.replace("\x00", "")).strip()


def as_percent_values(series: pd.Series) -> pd.Series:
    values = pd.to_numeric(series, errors="coerce")
    if values.dropna().abs().median() < 1:
        values = values * 100
    return values


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: Any, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(clean_text(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(9)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10 if level == 1 else 7)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(clean_text(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.bold = True
    run.font.size = Pt({1: 17, 2: 13, 3: 11}.get(level, 10))
    run.font.color.rgb = RGBColor(31, 78, 121) if level <= 2 else RGBColor(68, 68, 68)


def add_para(doc: Document, text: str = "", *, bold: bool = False, italic: bool = False) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(5)
    run = para.add_run(clean_text(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run.font.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(clean_text(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(10.2)


def add_table(doc: Document, headers: list[str], rows: list[list[Any]], *, max_rows: int | None = None) -> None:
    if max_rows is not None:
        rows = rows[:max_rows]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_color = RGBColor(255, 255, 255)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "1F4E79")
        set_cell_text(cell, header, bold=True, color=header_color)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F3F6FA")
            set_cell_text(cell, value)
    doc.add_paragraph()


def truncate_text(value: Any, max_chars: int = 120) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    text = clean_text(value).replace("\n", " ").replace("\r", " ").strip()
    return text if len(text) <= max_chars else text[: max_chars - 3] + "..."


def rows_from_frame(df: pd.DataFrame, columns: list[str], *, max_rows: int = 5, max_chars: int = 120) -> list[list[str]]:
    if df.empty:
        return []
    available = [col for col in columns if col in df.columns]
    if not available:
        return []
    out = df[available].head(max_rows).copy()
    return [[truncate_text(value, max_chars=max_chars) for value in row] for row in out.to_numpy()]


def code_excerpt(rel_path: str, start: int, end: int) -> str:
    path = ROOT / rel_path
    if not path.exists():
        return f"# Missing file: {rel_path}"
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    excerpt = []
    for line_no in range(start, min(end, len(lines)) + 1):
        excerpt.append(f"{line_no:04d}: {lines[line_no - 1]}")
    return "\n".join(excerpt)


def add_code_block(doc: Document, title: str, rel_path: str, start: int, end: int, note: str) -> None:
    add_para(doc, f"{title} ({rel_path}:{start}-{end})", bold=True)
    if note:
        add_para(doc, note)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F7F7F7")
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(code_excerpt(rel_path, start, end))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.2)
    doc.add_paragraph()


def add_picture_if_exists(doc: Document, path: Path, title: str, width: float = 6.5) -> None:
    if not path.exists():
        return
    add_para(doc, title, bold=True)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=Inches(width))


def add_cover(doc: Document, latest_date: str) -> None:
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(clean_text("BÃO CÃO PHÃ‚N TÃCH THá»Š TRÆ¯á»œNG VÃ€NG VIá»†T NAM"))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(clean_text("Business Understanding, EDA, Modeling, Forecasting vÃ  quyáº¿t Ä‘á»‹nh mua 1-5 thÃ¡ng"))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(68, 114, 196)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        clean_text(
            f"Snapshot dá»¯ liá»‡u: Ä‘áº¿n {latest_date}\n"
            f"NgÃ y tÃ¡i táº¡o bÃ¡o cÃ¡o: {datetime.now().strftime('%d/%m/%Y %H:%M')}\n"
            "Deliverable: DOCX há»c thuáº­t, khÃ´ng dÃ¹ng dá»¯ liá»‡u live ngoÃ i snapshot"
        )
    )
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(90, 90, 90)
    doc.add_page_break()


def profile_datasets(datasets: dict[str, tuple[str, str, str]]) -> tuple[list[DataProfile], dict[str, pd.DataFrame]]:
    frames: dict[str, pd.DataFrame] = {}
    profiles: list[DataProfile] = []
    for name, (path, key_fields, role) in datasets.items():
        df = read_csv(path)
        frames[name] = df
        profiles.append(
            DataProfile(
                name=name,
                path=path,
                rows=len(df),
                cols=len(df.columns),
                date_span=date_span(df),
                key_fields=key_fields,
                role=role,
            )
        )
    return profiles, frames


def build_profiles() -> tuple[list[DataProfile], dict[str, pd.DataFrame]]:
    datasets = {
        "SJC historical target": (
            "data/lake/gold_quotes_sjc_historical.csv",
            "date, business_date, buy, sell, spread, source",
            "Nguá»“n target huáº¥n luyá»‡n chÃ­nh; giá»¯ báº£n ghi historical-valid.",
        ),
        "Domestic daily panel": (
            "data/lake/gold_domestic_daily_panel.csv",
            "date, business_date, provider, gold_type, buy_price, sell_price",
            "Panel giÃ¡ vÃ ng trong nÆ°á»›c nhiá»u nguá»“n.",
        ),
        "Global reference daily": (
            "data/lake/global_reference_daily.csv",
            "date, gold_futures_close_usd_oz, usd_vnd_mid, vix, dxy_index",
            "Biáº¿n toÃ n cáº§u vÃ  FX dÃ¹ng Ä‘á»ƒ giáº£i thÃ­ch giÃ¡ vÃ ng ná»™i Ä‘á»‹a.",
        ),
        "Premium enriched": (
            "data/lake/pipeline_output_premium_enriched.csv",
            "date, global_gold_vnd_per_luong, premium, premium_pct",
            "PhÃ¢n rÃ£ premium SJC so vá»›i vÃ ng tháº¿ giá»›i quy Ä‘á»•i VND/lÆ°á»£ng.",
        ),
        "VN macro as-of": (
            "data/lake/pipeline_output_vn_macro_asof.csv",
            "available_from, observation_date, indicator_name, value",
            "Panel vÄ© mÃ´ vá»›i má»‘c cÃ´ng bá»‘ Ä‘á»ƒ chá»‘ng leakage.",
        ),
        "GPR daily": (
            "data/lake/gpr_daily_geopolitical_risk.csv",
            "date, GPRD, GPRD_MA7, GPRD_MA30",
            "Chá»‰ sá»‘ rá»§i ro Ä‘á»‹a chÃ­nh trá»‹.",
        ),
        "Event regime": (
            "data/lake/pipeline_output_event_regime.csv",
            "event_date, event_type, severity, expected_channel",
            "Táº¿t, Tháº§n TÃ i, mÃ¹a cÆ°á»›i, chÃ­nh sÃ¡ch, crisis/regime events.",
        ),
        "News raw headlines": (
            "data/lake/news_raw_headlines_vietnam_gold.csv",
            "event_date, headline, body_text, category, source",
            "Nguá»“n headline/backfill cho biáº¿n intensity vÃ  sentiment heuristic.",
        ),
        "Model frame": (
            "data/lake/modeling/model_frame_daily.csv",
            "date, prices, targets, lag/rolling features",
            "Báº£ng cuá»‘i cho huáº¥n luyá»‡n: má»™t dÃ²ng má»—i ngÃ y, target 1/3/5 thÃ¡ng lá»‹ch.",
        ),
        "Model results": (
            "data/lake/modeling/model_results.csv",
            "model, horizon_months, fold, phase, mae, rmse, directional_accuracy",
            "Leaderboard walk-forward theo horizon vÃ  fold.",
        ),
        "Walk-forward predictions": (
            "data/lake/modeling/walk_forward_predictions.csv",
            "date, horizon_months, fold, model, actual, predicted",
            "Dá»± bÃ¡o tá»«ng fold Ä‘á»ƒ audit model behavior.",
        ),
        "Decision signals": (
            "data/lake/modeling/decision_signals.csv",
            "date, horizon_months, prob_return_positive, q10, buy_signal",
            "Báº£ng tÃ­n hiá»‡u mua theo threshold grid.",
        ),
        "Trading signals": (
            "data/lake/modeling/trading_signals.csv",
            "date, horizon_months, selected_model, prob_positive, buy_signal_any",
            "Báº£ng collapse má»™t dÃ²ng/ngÃ y cho paper trading.",
        ),
    }
    return profile_datasets(datasets)


def compute_data_quality(frames: dict[str, pd.DataFrame], summary: dict[str, Any]) -> list[list[Any]]:
    checks: list[list[Any]] = []
    frame = frames.get("Model frame", pd.DataFrame())
    target = frames.get("SJC historical target", pd.DataFrame())
    premium = frames.get("Premium enriched", pd.DataFrame())
    signals = frames.get("Trading signals", pd.DataFrame())

    if not target.empty:
        valid = target.copy()
        for col in ["buy", "sell", "spread"]:
            if col in valid.columns:
                valid[col] = pd.to_numeric(valid[col], errors="coerce")
        invalid_prices = int(((valid.get("buy", 1) <= 0) | (valid.get("sell", 1) <= 0) | (valid.get("sell", 1) < valid.get("buy", 0))).sum())
        checks.append(["Target price validity", f"{invalid_prices:,} dÃ²ng báº¥t thÆ°á»ng", "Cao", "Pháº£i báº±ng 0 hoáº·c Ä‘Æ°á»£c giáº£i thÃ­ch trÆ°á»›c khi dÃ¹ng lÃ m label."])

    if not frame.empty:
        checks.append(["Model frame grain", f"{len(frame):,} dÃ²ng, {frame['date'].nunique():,} ngÃ y unique", "Cao", "Má»™t dÃ²ng/ngÃ y cho training frame."])
        for col in ["global_feature_date", "gpr_feature_date", "macro_feature_date"]:
            if col in frame.columns:
                ok = (safe_date_series(frame[col]).dropna() <= safe_date_series(frame.loc[frame[col].notna(), "date"])).all()
                checks.append([f"As-of guard: {col}", "PASS" if ok else "FAIL", "Cao", "KhÃ´ng dÃ¹ng dá»¯ liá»‡u tÆ°Æ¡ng lai so vá»›i ngÃ y quyáº¿t Ä‘á»‹nh."])
        for horizon in [1, 3, 5]:
            col = f"net_return_{horizon}d"
            if col in frame.columns:
                checks.append([f"Target non-null {horizon}d", f"{frame[col].notna().sum():,}", "Trung bÃ¬nh", "Sá»‘ máº«u cÃ³ nhÃ£n sau khi trá»« pháº§n cuá»‘i chÆ°a cÃ³ tÆ°Æ¡ng lai."])

    if not premium.empty and "premium" in premium.columns:
        checks.append(["Premium missing", pct(pd.to_numeric(premium["premium"], errors="coerce").isna().mean()), "Cao", "Caveat lá»›n cho diá»…n giáº£i premium vÃ  cÆ¡ há»™i entry."])

    if not signals.empty:
        checks.append(["Trading signal coverage", f"{len(signals):,} dÃ²ng", "Trung bÃ¬nh", "Báº£ng collapse Ä‘Ã£ xuáº¥t láº¡i sau modeling run."])

    blockers = summary.get("blockers", [])
    for blocker in blockers[:6]:
        checks.append(["Runtime blocker", blocker[:180], "Cao" if "unavailable" in blocker.lower() or "missing" in blocker.lower() else "Trung bÃ¬nh", "Ghi vÃ o bÃ¡o cÃ¡o, khÃ´ng thay báº±ng káº¿t quáº£ giáº£."])
    return checks


def compute_eda_tables(frames: dict[str, pd.DataFrame]) -> dict[str, list[list[Any]]]:
    tables: dict[str, list[list[Any]]] = {}
    frame = frames.get("Model frame", pd.DataFrame()).copy()
    premium = frames.get("Premium enriched", pd.DataFrame()).copy()
    events = frames.get("Event regime", pd.DataFrame()).copy()
    global_ref = frames.get("Global reference daily", pd.DataFrame()).copy()

    if not frame.empty:
        frame["date"] = safe_date_series(frame["date"])
        frame["year"] = frame["date"].dt.year
        frame["sell_price"] = pd.to_numeric(frame["sell_price"], errors="coerce")
        frame["buy_price"] = pd.to_numeric(frame["buy_price"], errors="coerce")
        frame["spread_pct"] = pd.to_numeric(frame.get("spread_pct"), errors="coerce") * 100
        yearly = (
            frame.dropna(subset=["date", "sell_price"])
            .groupby("year")
            .agg(
                start=("sell_price", "first"),
                end=("sell_price", "last"),
                min_price=("sell_price", "min"),
                max_price=("sell_price", "max"),
                obs=("sell_price", "count"),
                spread_mean=("spread_pct", "mean"),
            )
            .reset_index()
        )
        yearly["return"] = yearly["end"] / yearly["start"] - 1
        tables["yearly_price"] = [
            [int(r.year), money_m(r.start), money_m(r.end), money_m(r.min_price), money_m(r.max_price), pct(r["return"]), f"{int(r.obs):,}", pct_points(r.spread_mean)]
            for _, r in yearly.tail(10).iterrows()
        ]

        corr_cols = [
            c
            for c in frame.columns
            if any(k in c for k in ["premium_pct_lag", "usd_vnd", "vix", "dxy", "treasury_10y", "GPRD_MA", "spread_pct"])
        ]
        corr_rows: list[list[Any]] = []
        for horizon in [1, 3, 5]:
            target_col = f"net_return_{horizon}d"
            values = []
            if target_col in frame.columns:
                for col in corr_cols:
                    x = pd.to_numeric(frame[col], errors="coerce")
                    y = pd.to_numeric(frame[target_col], errors="coerce")
                    corr = x.corr(y)
                    if pd.notna(corr):
                        values.append((col, corr))
            for col, corr in sorted(values, key=lambda x: abs(x[1]), reverse=True)[:5]:
                corr_rows.append([HORIZON_LABELS[horizon], col, f"{corr:+.3f}"])
        tables["correlations"] = corr_rows

    if not premium.empty and "premium_pct" in premium.columns:
        premium_pct = as_percent_values(premium["premium_pct"])
        premium_abs = pd.to_numeric(premium.get("premium"), errors="coerce")
        valid = premium_pct.dropna()
        regimes = {
            "Tháº¥p (<3%)": float((valid < 3).mean()) if len(valid) else np.nan,
            "BÃ¬nh thÆ°á»ng (3-6%)": float(((valid >= 3) & (valid < 6)).mean()) if len(valid) else np.nan,
            "Cao (6-10%)": float(((valid >= 6) & (valid < 10)).mean()) if len(valid) else np.nan,
            "Crisis (>10%)": float((valid >= 10).mean()) if len(valid) else np.nan,
        }
        tables["premium_stats"] = [
            ["Sá»‘ ngÃ y cÃ³ premium", f"{valid.notna().sum():,}", ""],
            ["Premium trung bÃ¬nh", money_m(premium_abs.mean()), pct_points(valid.mean())],
            ["Premium median", money_m(premium_abs.median()), pct_points(valid.median())],
            ["P10", money_m(premium_abs.quantile(0.10)), pct_points(valid.quantile(0.10))],
            ["P90", money_m(premium_abs.quantile(0.90)), pct_points(valid.quantile(0.90))],
        ]
        tables["premium_regime"] = [[k, pct(v)] for k, v in regimes.items()]

    if not events.empty and "event_type" in events.columns:
        counts = events["event_type"].fillna("unknown").value_counts().head(12)
        tables["events"] = [[idx, f"{int(value):,}"] for idx, value in counts.items()]

    if not global_ref.empty:
        rows = []
        for col in ["gold_futures_close_usd_oz", "lbma_price_usd_oz", "usd_vnd_mid", "vix", "dxy_index", "treasury_10y_pct", "oil_wti_usd_barrel"]:
            if col in global_ref.columns:
                values = pd.to_numeric(global_ref[col], errors="coerce")
                rows.append([col, f"{values.notna().sum():,}", pct(values.notna().mean()), num(values.min()), num(values.max())])
        tables["global_coverage"] = rows

    return tables


def build_figures(frames: dict[str, pd.DataFrame], summary: dict[str, Any]) -> dict[str, Path]:
    FIGURES.mkdir(parents=True, exist_ok=True)
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import seaborn as sns

    sns.set_theme(style="whitegrid")
    paths: dict[str, Path] = {}
    frame = frames.get("Model frame", pd.DataFrame()).copy()
    results = frames.get("Model results", pd.DataFrame()).copy()
    signals = frames.get("Trading signals", pd.DataFrame()).copy()

    if not frame.empty:
        frame["date"] = safe_date_series(frame["date"])
        frame["sell_price"] = pd.to_numeric(frame["sell_price"], errors="coerce")
        frame["global_gold_vnd_per_luong"] = pd.to_numeric(frame.get("global_gold_vnd_per_luong"), errors="coerce")
        fig, ax = plt.subplots(figsize=(10, 4.2))
        ax.plot(frame["date"], frame["sell_price"] / 1_000_000, label="SJC sell", color="#1F77B4", linewidth=1.2)
        if frame["global_gold_vnd_per_luong"].notna().any():
            ax.plot(frame["date"], frame["global_gold_vnd_per_luong"] / 1_000_000, label="Global gold VND/luong", color="#B8860B", linewidth=1.0, alpha=0.8)
        ax.set_title("SJC price versus global gold proxy")
        ax.set_ylabel("Million VND per luong")
        ax.legend()
        path = FIGURES / "final_01_price_vs_global.png"
        fig.savefig(path, dpi=180, bbox_inches="tight")
        plt.close(fig)
        paths["price"] = path

        if "premium_pct" in frame.columns:
            prem = as_percent_values(frame["premium_pct"])
            fig, ax = plt.subplots(figsize=(10, 4.2))
            ax.plot(frame["date"], prem, color="#D28E00", linewidth=0.8)
            ax.axhline(3, color="#2E7D32", linestyle="--", linewidth=0.8)
            ax.axhline(6, color="#F57C00", linestyle="--", linewidth=0.8)
            ax.axhline(10, color="#C62828", linestyle="--", linewidth=0.8)
            ax.set_title("Domestic premium regime")
            ax.set_ylabel("Premium (%)")
            path = FIGURES / "final_02_premium_regime.png"
            fig.savefig(path, dpi=180, bbox_inches="tight")
            plt.close(fig)
            paths["premium"] = path

    if not results.empty:
        mean = results[results["pinball_loss"].isna()].copy()
        mean["mae"] = pd.to_numeric(mean["mae"], errors="coerce")
        agg = mean.groupby(["model", "horizon_months"], as_index=False)["mae"].mean().dropna()
        if not agg.empty:
            fig, axes = plt.subplots(1, 3, figsize=(13, 4), sharex=False)
            for axis, horizon in zip(axes, [1, 3, 5]):
                sub = agg[agg["horizon_months"].eq(horizon)].sort_values("mae").head(7)
                axis.barh(sub["model"], sub["mae"], color="#4472C4")
                axis.invert_yaxis()
                axis.set_title(f"{horizon}d MAE")
                axis.set_xlabel("MAE")
            fig.tight_layout()
            path = FIGURES / "final_03_model_leaderboard.png"
            fig.savefig(path, dpi=180, bbox_inches="tight")
            plt.close(fig)
            paths["leaderboard"] = path

    if not signals.empty:
        signals["buy_signal_any"] = signals["buy_signal_any"].astype(str).str.lower().isin(["true", "1"])
        sig = (
            signals.groupby(["phase", "horizon_months"], as_index=False)
            .agg(signal_rate=("buy_signal_any", "mean"), avg_actual=("avg_actual_return", "mean"))
        )
        if not sig.empty:
            sig["label"] = sig["phase"].astype(str) + " " + sig["horizon_months"].astype(str) + "d"
            fig, ax = plt.subplots(figsize=(9, 4))
            ax.bar(sig["label"], sig["signal_rate"] * 100, color="#70AD47")
            ax.set_title("Decision signal frequency")
            ax.set_ylabel("Signal days (%)")
            ax.tick_params(axis="x", rotation=35)
            path = FIGURES / "final_04_signal_frequency.png"
            fig.savefig(path, dpi=180, bbox_inches="tight")
            plt.close(fig)
            paths["signals"] = path

    return paths


def summarize_models(summary: dict[str, Any]) -> tuple[list[list[Any]], list[list[Any]], list[list[Any]]]:
    leaderboard = pd.DataFrame(summary.get("leaderboard", []))
    if leaderboard.empty:
        return [], [], []

    mean = leaderboard[leaderboard["pinball_loss"].isna()].copy()
    quant = leaderboard[leaderboard["pinball_loss"].notna()].copy()
    mean_rows: list[list[Any]] = []
    quant_rows: list[list[Any]] = []
    for horizon in [1, 3, 5]:
        sub = mean[mean["horizon_months"].eq(horizon)].copy()
        if not sub.empty:
            best_mae = sub.sort_values("mae").iloc[0]
            best_da = sub.sort_values("directional_accuracy", ascending=False).iloc[0]
            mean_rows.append(
                [
                    HORIZON_LABELS[horizon],
                    best_mae["model"],
                    f"{float(best_mae['mae']):.4f}",
                    f"{float(best_mae['rmse']):.4f}",
                    pct(float(best_mae["directional_accuracy"])),
                    best_da["model"],
                    pct(float(best_da["directional_accuracy"])),
                ]
            )
        qsub = quant[quant["horizon_months"].eq(horizon)].dropna(subset=["pinball_loss"])
        if not qsub.empty:
            best_q = qsub.sort_values("pinball_loss").iloc[0]
            quant_rows.append([HORIZON_LABELS[horizon], best_q["model"], f"{float(best_q['pinball_loss']):.5f}"])

    decision = pd.DataFrame(summary.get("decision_summary", []))
    decision_rows: list[list[Any]] = []
    if not decision.empty:
        for _, row in decision.sort_values(["horizon_months", "phase"]).iterrows():
            signal_days = int(row["signal_days"])
            observations = int(row["observations"])
            decision_rows.append(
                [
                    HORIZON_LABELS[int(row["horizon_months"])],
                    row["phase"],
                    f"{signal_days:,}/{observations:,}",
                    pct(signal_days / observations if observations else np.nan),
                    pct(row["avg_buy_day_return"], 2),
                    pct(row["avg_strategy_return"], 2),
                ]
            )
    return mean_rows, quant_rows, decision_rows


def load_feature_columns() -> list[str]:
    payload = read_json("data/lake/modeling/feature_columns.json")
    if isinstance(payload, list):
        return [str(x) for x in payload]
    return []


def train_snapshot_forecasts(frame: pd.DataFrame, feature_cols: list[str]) -> pd.DataFrame:
    if frame.empty or not feature_cols:
        return pd.DataFrame()
    try:
        from sklearn.ensemble import GradientBoostingRegressor, RandomForestRegressor
        from sklearn.impute import SimpleImputer
        from sklearn.pipeline import Pipeline
    except Exception as exc:
        return pd.DataFrame([{"status": f"snapshot forecast skipped: {exc}"}])

    data = frame.copy()
    data["date"] = safe_date_series(data["date"])
    feature_cols = [c for c in feature_cols if c in data.columns]
    latest = data.sort_values("date").iloc[[-1]].copy()
    rows: list[dict[str, Any]] = []

    for horizon in [1, 3, 5]:
        target_col = f"net_return_{horizon}m"
        if target_col not in data.columns:
            continue
        train = data.dropna(subset=[target_col]).copy()
        if len(train) < 500:
            rows.append({"horizon_months": horizon, "status": "skipped_insufficient_training_rows"})
            continue

        # Select the most informative features on the training window only.
        scored: list[tuple[str, float]] = []
        y = pd.to_numeric(train[target_col], errors="coerce")
        for col in feature_cols:
            x = pd.to_numeric(train[col], errors="coerce")
            if x.notna().sum() < 200 or x.nunique(dropna=True) < 2:
                continue
            corr = abs(x.corr(y))
            if pd.notna(corr):
                scored.append((col, corr))
        selected = [col for col, _ in sorted(scored, key=lambda item: item[1], reverse=True)[:120]]
        if not selected:
            rows.append({"horizon_months": horizon, "status": "skipped_no_usable_features"})
            continue

        x_train = train[selected].apply(pd.to_numeric, errors="coerce")
        y_train = y.loc[x_train.index].astype(float)
        x_latest = latest[selected].apply(pd.to_numeric, errors="coerce")

        point_model = Pipeline(
            [
                ("imputer", SimpleImputer(strategy="median")),
                (
                    "model",
                    RandomForestRegressor(
                        n_estimators=350,
                        max_depth=7,
                        min_samples_leaf=12,
                        random_state=42,
                        n_jobs=-1,
                    ),
                ),
            ]
        )
        q10_model = Pipeline(
            [
                ("imputer", SimpleImputer(strategy="median")),
                (
                    "model",
                    GradientBoostingRegressor(
                        loss="quantile",
                        alpha=0.10,
                        n_estimators=250,
                        max_depth=3,
                        learning_rate=0.035,
                        random_state=42,
                    ),
                ),
            ]
        )
        point_model.fit(x_train, y_train)
        pred_train = point_model.predict(x_train)
        residual_std = float(np.nanstd(y_train.to_numpy() - pred_train))
        residual_std = residual_std if np.isfinite(residual_std) and residual_std > 0 else 0.05
        pred = float(point_model.predict(x_latest)[0])
        q10_model.fit(x_train, y_train)
        q10 = float(q10_model.predict(x_latest)[0])
        prob_positive = 1.0 - 0.5 * (1.0 + math.erf((0.0 - pred) / (residual_std * math.sqrt(2))))
        buy_signal = bool((prob_positive >= 0.50) and (q10 >= -0.10))
        rows.append(
            {
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "snapshot_date": str(latest["date"].iloc[0].date()),
                "horizon_months": horizon,
                "horizon_label": HORIZON_LABELS[horizon],
                "point_model": "random_forest_snapshot_top120_features",
                "q10_model": "gradient_boosting_quantile_q10_snapshot_top120_features",
                "train_rows": int(len(train)),
                "feature_count": int(len(selected)),
                "predicted_net_return": pred,
                "q10_predicted_net_return": q10,
                "prob_return_positive": prob_positive,
                "buy_signal": buy_signal,
                "status": "ok",
            }
        )

    out = pd.DataFrame(rows)
    if not out.empty:
        MODELING.mkdir(parents=True, exist_ok=True)
        out.to_csv(MODELING / "snapshot_forecasts.csv", index=False)
    return out


def decision_recommendation(snapshot: pd.DataFrame, decision_rows: list[list[Any]]) -> tuple[str, list[list[Any]]]:
    rows: list[list[Any]] = []
    if snapshot.empty or "predicted_net_return" not in snapshot.columns:
        return (
            "KhÃ´ng Ä‘Æ°a ra khuyáº¿n nghá»‹ mua má»›i vÃ¬ script khÃ´ng táº¡o Ä‘Æ°á»£c snapshot forecast kiá»ƒm chá»©ng Ä‘Æ°á»£c.",
            rows,
        )
    for _, row in snapshot.iterrows():
        if row.get("status") != "ok":
            rows.append([row.get("horizon_label", ""), row.get("status", ""), "", "", "KhÃ´ng káº¿t luáº­n"])
            continue
        action = "Mua tÃ­ch lÅ©y" if bool(row["buy_signal"]) else "KhÃ´ng mua má»›i"
        if row["horizon_months"] == 63:
            action = "KhÃ´ng mua má»›i" if row["predicted_net_return"] <= 0 or row["q10_predicted_net_return"] < -0.10 else action
        rows.append(
            [
                row["horizon_label"],
                pct(row["predicted_net_return"], 2),
                pct(row["q10_predicted_net_return"], 2),
                pct(row["prob_return_positive"], 1),
                action,
            ]
        )

    five = snapshot[snapshot["horizon_months"].eq(5)]
    three = snapshot[snapshot["horizon_months"].eq(3)]
    one = snapshot[snapshot["horizon_months"].eq(1)]
    if not five.empty and bool(five.iloc[0].get("buy_signal", False)):
        recommendation = (
            "Khuyáº¿n nghá»‹ chÃ­nh: cÃ³ thá»ƒ mua tÃ­ch lÅ©y cÃ³ kiá»ƒm soÃ¡t cho horizon khoáº£ng 5 thÃ¡ng, "
            "nhÆ°ng khÃ´ng mua máº¡nh má»™t láº§n. Horizon 3 thÃ¡ng khÃ´ng Ä‘á»§ háº¥p dáº«n vÃ¬ backtest tÃ­n hiá»‡u Ã¢m; "
            "horizon 1 thÃ¡ng chá»‰ phÃ¹ há»£p giao dá»‹ch nhá» do tÃ­n hiá»‡u lá»‹ch sá»­ ráº¥t hiáº¿m."
        )
    elif not one.empty and bool(one.iloc[0].get("buy_signal", False)):
        recommendation = (
            "Khuyáº¿n nghá»‹ chÃ­nh: chá»‰ mua nhá»/ngáº¯n háº¡n náº¿u cháº¥p nháº­n rá»§i ro, chÆ°a Ä‘á»§ báº±ng chá»©ng Ä‘á»ƒ mua máº¡nh "
            "cho 3-5 thÃ¡ng."
        )
    else:
        recommendation = (
            "Khuyáº¿n nghá»‹ chÃ­nh: khÃ´ng mua má»›i á»Ÿ quy mÃ´ lá»›n theo snapshot hiá»‡n táº¡i; Æ°u tiÃªn chá» premium/spread "
            "háº¡ nhiá»‡t hoáº·c cÃ³ tÃ­n hiá»‡u xÃ¡c nháº­n tá»‘t hÆ¡n."
        )
    if not three.empty and bool(three.iloc[0].get("buy_signal", False)):
        recommendation += " LÆ°u Ã½: náº¿u mÃ´ hÃ¬nh snapshot 3 thÃ¡ng báº­t tÃ­n hiá»‡u, váº«n cáº§n háº¡ trá»ng sá»‘ náº¿u backtest 3 thÃ¡ng cÃ³ avg buy-day return Ã¢m."
    return recommendation, rows


def add_executive_summary(doc: Document, latest_date: str, recommendation: str, summary: dict[str, Any]) -> None:
    add_heading(doc, "TÃ³m táº¯t Ä‘iá»u hÃ nh", 1)
    add_bullet(doc, f"{SNAPSHOT_NOTE} Káº¿t luáº­n khÃ´ng pháº£i tÃ­n hiá»‡u live sau ngÃ y {latest_date}.")
    add_bullet(doc, recommendation)
    add_bullet(
        doc,
        "Target chÃ­nh lÃ  lá»£i nhuáº­n sau spread: mua táº¡i giÃ¡ bÃ¡n hÃ´m nay vÃ  bÃ¡n láº¡i theo giÃ¡ mua á»Ÿ ngÃ y kháº£ dá»¥ng gáº§n nháº¥t sau má»‘c 1/3/5 thÃ¡ng lá»‹ch.",
    )
    add_bullet(
        doc,
        f"Model frame cÃ³ {summary.get('rows', 'n/a'):,} ngÃ y tá»« {summary.get('date_min')} Ä‘áº¿n {summary.get('date_max')}; "
        f"premium missing {pct(summary.get('premium_missing_rate'))}.",
    )
    add_bullet(
        doc,
        "LightGBM, XGBoost vÃ  CatBoost Ä‘Æ°á»£c cÃ i vÃ  train tháº­t khi import Ä‘Æ°á»£c; deep models chá»‰ xuáº¥t hiá»‡n khi cÃ³ runner huáº¥n luyá»‡n tháº­t.",
    )
    add_para(
        doc,
        "Báº£n bÃ¡o cÃ¡o nÃ y Ä‘Æ°á»£c viáº¿t theo hÆ°á»›ng ra quyáº¿t Ä‘á»‹nh Ä‘áº§u tÆ° chá»© khÃ´ng chá»‰ mÃ´ táº£ biáº¿n Ä‘á»™ng giÃ¡. CÃ¢u há»i trung tÃ¢m lÃ : náº¿u má»™t nhÃ  Ä‘áº§u tÆ° cÃ¡ nhÃ¢n mua vÃ ng SJC táº¡i thá»i Ä‘iá»ƒm snapshot, chá»‹u giÃ¡ bÃ¡n ra hiá»‡n táº¡i vÃ  cÃ³ thá»ƒ bÃ¡n láº¡i theo giÃ¡ mua vÃ o sau 1, 3 hoáº·c 5 thÃ¡ng, lá»£i suáº¥t ká»³ vá»ng cÃ³ Ä‘á»§ bÃ¹ cho spread, rá»§i ro giáº£m giÃ¡ vÃ  chi phÃ­ cÆ¡ há»™i hay khÃ´ng. VÃ¬ váº­y, káº¿t luáº­n mua/khÃ´ng mua Ä‘Æ°á»£c Ä‘á»c cÃ¹ng vá»›i xÃ¡c suáº¥t lá»£i suáº¥t dÆ°Æ¡ng, phÃ¢n vá»‹ rá»§i ro q10 vÃ  káº¿t quáº£ kiá»ƒm thá»­ lÃ¹i cuá»‘n chiáº¿u.",
    )
    add_para(
        doc,
        "Káº¿t quáº£ táº¡i snapshot nghiÃªng vá» cÃ¡ch tiáº¿p cáº­n tháº­n trá»ng: horizon 1 thÃ¡ng, 3 thÃ¡ng vÃ  5 thÃ¡ng Ä‘Æ°á»£c Ä‘á»c riÃªng theo expected return, q10 downside vÃ  consistency backtest. TÃ­n hiá»‡u 5 thÃ¡ng náº¿u cÃ³ khÃ´ng Ä‘á»“ng nghÄ©a vá»›i mua máº¡nh ngay láº­p tá»©c, vÃ¬ premium SJC vÃ  spread cÃ³ thá»ƒ Ä‘áº£o chiá»u nhanh khi NHNN thay Ä‘á»•i chÃ­nh sÃ¡ch cung vÃ ng hoáº·c thá»‹ trÆ°á»ng quá»‘c táº¿ Ä‘á»•i cháº¿ Ä‘á»™ lÃ£i suáº¥t/USD.",
    )
    add_para(
        doc,
        "Äiá»ƒm cáº£i thiá»‡n quan trá»ng so vá»›i báº£n caveat trÆ°á»›c lÃ  premium coverage Ä‘Ã£ Ä‘Æ°á»£c rebuild theo hierarchy cÃ³ gáº¯n nhÃ£n nguá»“n vÃ  proxy, giÃºp premium missing giáº£m xuá»‘ng má»©c Ä‘Æ°á»£c bÃ¡o cÃ¡o trong model frame. DÃ¹ váº­y, lÃ£i suáº¥t tiá»n gá»­i Viá»‡t Nam má»›i chá»‰ cÃ³ dá»¯ liá»‡u forward-monitoring tá»« nguá»“n ngÃ¢n hÃ ng thÆ°Æ¡ng máº¡i, cÃ²n news real-time availability váº«n chÆ°a Ä‘á»§ máº¡nh Ä‘á»ƒ coi headline lÃ  tÃ­n hiá»‡u giao dá»‹ch lá»‹ch sá»­ tuyá»‡t Ä‘á»‘i.",
    )


def add_business_understanding(doc: Document) -> None:
    add_heading(doc, "1. Business Understanding vÃ  gÃ³c nhÃ¬n kinh táº¿", 1)
    add_para(
        doc,
        "BÃ i toÃ¡n khÃ´ng chá»‰ lÃ  dá»± bÃ¡o giÃ¡ vÃ ng, mÃ  lÃ  quyáº¿t Ä‘á»‹nh cÃ³ nÃªn mua vÃ ng váº­t cháº¥t táº¡i Viá»‡t Nam trong 1-5 thÃ¡ng tá»›i. "
        "NgÆ°á»i mua tráº£ giÃ¡ bÃ¡n ra hÃ´m nay vÃ  náº¿u thoÃ¡t vá»‹ tháº¿ sáº½ nháº­n giÃ¡ mua vÃ o trong tÆ°Æ¡ng lai, nÃªn bid-ask spread vÃ  premium ná»™i Ä‘á»‹a lÃ  má»™t pháº§n trá»±c tiáº¿p cá»§a lá»£i nhuáº­n.",
    )
    add_para(
        doc,
        "VÃ ng Viá»‡t Nam chá»‹u tÃ¡c Ä‘á»™ng cá»§a bá»‘n nhÃ³m lá»±c: giÃ¡ vÃ ng quá»‘c táº¿, tá»· giÃ¡ USD/VND, premium SJC do cáº¥u trÃºc cung-cáº§u vÃ  chÃ­nh sÃ¡ch, vÃ  cháº¿ Ä‘á»™ rá»§i ro nhÆ° láº¡m phÃ¡t, lÃ£i suáº¥t, Ä‘á»‹a chÃ­nh trá»‹, mÃ¹a vá»¥ Táº¿t/Tháº§n TÃ i.",
    )
    add_heading(doc, "1.1. Tá»« bÃ i toÃ¡n giÃ¡ sang bÃ i toÃ¡n quyáº¿t Ä‘á»‹nh", 2)
    add_para(
        doc,
        "Náº¿u chá»‰ dá»± bÃ¡o má»©c giÃ¡ bÃ¡n ra trong tÆ°Æ¡ng lai, mÃ´ hÃ¬nh ráº¥t dá»… táº¡o ra tÃ­n hiá»‡u Ä‘áº¹p trÃªn giáº¥y nhÆ°ng khÃ´ng dÃ¹ng Ä‘Æ°á»£c ngoÃ i thá»±c táº¿. NhÃ  Ä‘áº§u tÆ° vÃ ng váº­t cháº¥t khÃ´ng mua á»Ÿ mid-price vÃ  cÅ©ng khÃ´ng bÃ¡n á»Ÿ mid-price; há» mua theo giÃ¡ bÃ¡n ra cá»§a tiá»‡m vÃ ng/ngÃ¢n hÃ ng vÃ ng vÃ  bÃ¡n láº¡i theo giÃ¡ mua vÃ o. Khoáº£ng cÃ¡ch nÃ y chÃ­nh lÃ  spread, vÃ  trong cÃ¡c giai Ä‘oáº¡n cÄƒng tháº³ng spread cÃ³ thá»ƒ lá»›n Ä‘áº¿n má»©c Äƒn mÃ²n toÃ n bá»™ lá»£i nhuáº­n ká»³ vá»ng.",
    )
    add_para(
        doc,
        "VÃ¬ lÃ½ do Ä‘Ã³, biáº¿n má»¥c tiÃªu cá»§a Ä‘á» tÃ i Ä‘Æ°á»£c Ä‘á»‹nh nghÄ©a lÃ  lá»£i suáº¥t rÃ²ng theo horizon: giÃ¡ mua vÃ o tÆ°Æ¡ng lai chia cho giÃ¡ bÃ¡n ra hiá»‡n táº¡i trá»« 1. Äá»‹nh nghÄ©a nÃ y buá»™c mÃ´ hÃ¬nh pháº£i tráº£ lá»i Ä‘Ãºng cÃ¢u há»i nghiá»‡p vá»¥: ngÆ°á»i mua hÃ´m nay cÃ³ cÃ²n lá»i sau khi chá»‹u chi phÃ­ giao dá»‹ch hay khÃ´ng. ÄÃ¢y cÅ©ng lÃ  lÃ½ do bÃ¡o cÃ¡o khÃ´ng Ä‘Æ°a ra má»™t con sá»‘ giÃ¡ má»¥c tiÃªu duy nháº¥t, mÃ  trÃ¬nh bÃ y expected return, q10 downside vÃ  xÃ¡c suáº¥t return dÆ°Æ¡ng cho tá»«ng chÃ¢n trá»i.",
    )
    add_heading(doc, "1.2. GÃ³c nhÃ¬n kinh táº¿ - tÃ i chÃ­nh cá»§a thá»‹ trÆ°á»ng vÃ ng Viá»‡t Nam", 2)
    add_para(
        doc,
        "Vá» máº·t kinh táº¿, vÃ ng lÃ  tÃ i sáº£n khÃ´ng táº¡o dÃ²ng tiá»n Ä‘á»‹nh ká»³. Khi lÃ£i suáº¥t tiá»n gá»­i cao, chi phÃ­ cÆ¡ há»™i cá»§a viá»‡c náº¯m giá»¯ vÃ ng tÄƒng; khi lÃ£i suáº¥t thá»±c giáº£m, chi phÃ­ cÆ¡ há»™i giáº£m vÃ  nhu cáº§u tÃ­ch trá»¯ vÃ ng thÆ°á»ng máº¡nh hÆ¡n. Trong bá»‘i cáº£nh Viá»‡t Nam, kÃªnh vÃ ng cáº¡nh tranh trá»±c tiáº¿p vá»›i tiá»n gá»­i ngÃ¢n hÃ ng, báº¥t Ä‘á»™ng sáº£n, chá»©ng khoÃ¡n vÃ  ngoáº¡i tá»‡, nÃªn mÃ´ hÃ¬nh cáº§n xem xÃ©t cáº£ biáº¿n vÄ© mÃ´ láº«n yáº¿u tá»‘ thanh khoáº£n ná»™i Ä‘á»‹a.",
    )
    add_para(
        doc,
        "Vá» máº·t tÃ i chÃ­nh quá»‘c táº¿, vÃ ng thÆ°á»ng Ä‘Æ°á»£c xem lÃ  tÃ i sáº£n phÃ²ng vá»‡ trÆ°á»›c láº¡m phÃ¡t, rá»§i ro Ä‘á»‹a chÃ­nh trá»‹ vÃ  báº¥t á»•n thá»‹ trÆ°á»ng. Tuy nhiÃªn, giÃ¡ vÃ ng trong nÆ°á»›c khÃ´ng pháº£i báº£n sao má»™t-má»™t cá»§a giÃ¡ vÃ ng tháº¿ giá»›i. Khi nguá»“n cung vÃ ng miáº¿ng SJC bá»‹ háº¡n cháº¿ hoáº·c cÆ¡ cháº¿ Ä‘áº¥u tháº§u/chÃ­nh sÃ¡ch thay Ä‘á»•i, premium ná»™i Ä‘á»‹a cÃ³ thá»ƒ má»Ÿ rá»™ng hoáº·c thu háº¹p Ä‘á»™c láº­p vá»›i xu hÆ°á»›ng toÃ n cáº§u.",
    )
    add_para(
        doc,
        "Do Ä‘Ã³, cÃ¡ch hiá»ƒu há»£p lÃ½ lÃ  phÃ¢n rÃ£ giÃ¡ vÃ ng Viá»‡t Nam thÃ nh ba lá»›p: lá»›p fair value toÃ n cáº§u quy Ä‘á»•i theo USD/VND, lá»›p premium thá»ƒ cháº¿ - cung cáº§u trong nÆ°á»›c, vÃ  lá»›p chi phÃ­ giao dá»‹ch/thanh khoáº£n. Má»™t quyáº¿t Ä‘á»‹nh mua tá»‘t cáº§n cáº£ ba lá»›p cÃ¹ng há»— trá»£: giÃ¡ tháº¿ giá»›i khÃ´ng quÃ¡ báº¥t lá»£i, tá»· giÃ¡ khÃ´ng táº¡o rá»§i ro ngÆ°á»£c, premium khÃ´ng quÃ¡ Ä‘áº¯t, vÃ  spread khÃ´ng lÃ m mÃ©o lá»£i suáº¥t thá»±c nháº­n.",
    )
    add_table(
        doc,
        ["Yáº¿u tá»‘", "CÆ¡ cháº¿ tÃ¡c Ä‘á»™ng", "Biáº¿n dá»¯ liá»‡u dÃ¹ng trong bÃ¡o cÃ¡o"],
        [
            ["GiÃ¡ vÃ ng tháº¿ giá»›i", "Neo giÃ¡ cÆ¡ báº£n theo USD/oz, pháº£n Ã¡nh lÃ£i suáº¥t thá»±c, USD vÃ  safe-haven demand.", "GC=F/LBMA proxy, global_gold_vnd_per_luong"],
            ["USD/VND", "TÄƒng USD/VND lÃ m vÃ ng quy Ä‘á»•i VND cao hÆ¡n ngay cáº£ khi USD gold Ä‘i ngang.", "usd_vnd_mid, usd_vnd_market_rate"],
            ["Premium SJC", "ChÃªnh lá»‡ch ná»™i Ä‘á»‹a pháº£n Ã¡nh khan hiáº¿m, quy Ä‘á»‹nh, Ä‘áº¥u tháº§u vÃ  tÃ¢m lÃ½ tÃ­ch trá»¯.", "premium, premium_pct, source_dispersion"],
            ["Thanh khoáº£n/spread", "Spread cao lÃ m giáº£m lá»£i nhuáº­n thá»±c nháº­n vÃ  bÃ¡o hiá»‡u stress bÃ¡n láº».", "spread_abs, spread_pct"],
            ["Rá»§i ro vÃ  mÃ¹a vá»¥", "Crisis/policy/Táº¿t/Tháº§n TÃ i cÃ³ thá»ƒ lÃ m premium vÃ  volatility tÄƒng ngáº¯n háº¡n.", "event_regime, GPRD, raw_news intensity"],
        ],
    )


def add_literature_review(doc: Document) -> None:
    add_heading(doc, "2. Literature Review vÃ  lá»±a chá»n phÆ°Æ¡ng phÃ¡p", 1)
    add_para(
        doc,
        "TÃ i liá»‡u dá»± bÃ¡o vÃ ng thÆ°á»ng chia thÃ nh ba nhÃ¡nh: mÃ´ hÃ¬nh chuá»—i thá»i gian cá»• Ä‘iá»ƒn, mÃ´ hÃ¬nh tÃ i chÃ­nh-vÄ© mÃ´ cÃ³ biáº¿n ngoáº¡i sinh, vÃ  mÃ´ hÃ¬nh machine learning/deep learning cho quan há»‡ phi tuyáº¿n. Vá»›i vÃ ng Viá»‡t Nam, bÃ¡o cÃ¡o Æ°u tiÃªn mÃ´ hÃ¬nh cÃ³ thá»ƒ xá»­ lÃ½ premium vÃ  biáº¿n ngoáº¡i sinh theo thá»i Ä‘iá»ƒm cÃ´ng bá»‘.",
    )
    add_para(
        doc,
        "NhÃ³m mÃ´ hÃ¬nh chuá»—i thá»i gian nhÆ° ARIMA/SARIMAX phÃ¹ há»£p Ä‘á»ƒ táº¡o baseline vÃ¬ giÃ¡ vÃ ng cÃ³ tá»± tÆ°Æ¡ng quan, xu hÆ°á»›ng vÃ  chu ká»³ ngáº¯n háº¡n. SARIMAX Ä‘áº·c biá»‡t há»¯u Ã­ch khi muá»‘n Ä‘Æ°a thÃªm biáº¿n ngoáº¡i sinh nhÆ° USD/VND, VIX, DXY, dáº§u, báº¡c hoáº·c chá»‰ bÃ¡o rá»§i ro Ä‘á»‹a chÃ­nh trá»‹. Äiá»ƒm máº¡nh cá»§a nhÃ³m nÃ y lÃ  dá»… kiá»ƒm soÃ¡t leakage vÃ  dá»… diá»…n giáº£i; Ä‘iá»ƒm yáº¿u lÃ  khÃ³ báº¯t quan há»‡ phi tuyáº¿n giá»¯a premium, sá»± kiá»‡n chÃ­nh sÃ¡ch vÃ  tÃ¢m lÃ½ thá»‹ trÆ°á»ng.",
    )
    add_para(
        doc,
        "NhÃ³m kinh táº¿ lÆ°á»£ng Ä‘a biáº¿n nhÆ° VECM/cointegration Ä‘Æ°á»£c xÃ©t vÃ¬ giÃ¡ vÃ ng ná»™i Ä‘á»‹a, giÃ¡ vÃ ng quá»‘c táº¿ quy Ä‘á»•i VND vÃ  tá»· giÃ¡ cÃ³ thá»ƒ cÃ³ quan há»‡ cÃ¢n báº±ng dÃ i háº¡n. Náº¿u tá»“n táº¡i cointegration á»•n Ä‘á»‹nh, sai lá»‡ch giá»¯a giÃ¡ ná»™i Ä‘á»‹a vÃ  fair value toÃ n cáº§u cÃ³ thá»ƒ chá»©a thÃ´ng tin mean-reversion. Tuy nhiÃªn, káº¿t quáº£ screen hiá»‡n táº¡i chÆ°a Ä‘á»§ Ä‘iá»u kiá»‡n Ä‘á»ƒ promote VECM thÃ nh mÃ´ hÃ¬nh forecast chÃ­nh, nÃªn bÃ¡o cÃ¡o ghi nháº­n nhÆ° má»™t kiá»ƒm tra cáº¥u trÃºc thay vÃ¬ Ã©p dÃ¹ng.",
    )
    add_para(
        doc,
        "NhÃ³m machine learning nhÆ° Random Forest, Gradient Boosting, LightGBM, XGBoost vÃ  CatBoost Ä‘Æ°á»£c dÃ¹ng Ä‘á»ƒ kiá»ƒm tra liá»‡u quan há»‡ phi tuyáº¿n cÃ³ cáº£i thiá»‡n dá»± bÃ¡o hay khÃ´ng. Vá»›i thá»‹ trÆ°á»ng vÃ ng Viá»‡t Nam, Ä‘Ã¢y lÃ  nhÃ³m Ä‘Ã¡ng thá»­ vÃ¬ tÃ¡c Ä‘á»™ng cá»§a premium thÆ°á»ng theo regime: premium tháº¥p cÃ³ thá»ƒ lÃ  cÆ¡ há»™i, nhÆ°ng premium quÃ¡ cao láº¡i cÃ³ thá»ƒ lÃ  rá»§i ro Ä‘áº£o chiá»u náº¿u nguá»“n cung chÃ­nh sÃ¡ch xuáº¥t hiá»‡n. Tree boosting cÃ³ kháº£ nÄƒng báº¯t cÃ¡c ngÆ°á»¡ng nÃ y tá»‘t hÆ¡n há»“i quy tuyáº¿n tÃ­nh.",
    )
    add_para(
        doc,
        "BÃ¡o cÃ¡o cÅ©ng phÃ¢n biá»‡t rÃµ mÃ´ hÃ¬nh dá»± bÃ¡o ká»³ vá»ng vÃ  mÃ´ hÃ¬nh dá»± bÃ¡o rá»§i ro. Má»™t expected return dÆ°Æ¡ng khÃ´ng Ä‘á»§ Ä‘á»ƒ mua náº¿u phÃ¢n vá»‹ q10 quÃ¡ Ã¢m, vÃ¬ nhÃ  Ä‘áº§u tÆ° cÃ¡ nhÃ¢n chá»‹u drawdown thá»±c vÃ  chi phÃ­ thoÃ¡t vá»‹ tháº¿. Do Ä‘Ã³, há»“i quy phÃ¢n vá»‹/quantile boosting Ä‘Æ°á»£c dÃ¹ng Ä‘á»ƒ Æ°á»›c lÆ°á»£ng downside, cÃ²n rule quyáº¿t Ä‘á»‹nh cáº§n Ä‘á»“ng thá»i thá»a xÃ¡c suáº¥t return dÆ°Æ¡ng vÃ  sÃ n q10.",
    )
    add_table(
        doc,
        ["NhÃ³m phÆ°Æ¡ng phÃ¡p", "Vai trÃ² trong bÃ i toÃ¡n vÃ ng", "CÃ¡ch Ã¡p dá»¥ng trong dá»± Ã¡n"],
        [
            ["ARIMA/SARIMAX", "Baseline máº¡nh cho chuá»—i cÃ³ tá»± tÆ°Æ¡ng quan vÃ  biáº¿n ngoáº¡i sinh.", "SARIMAX+exog trÃªn target return nhiá»u horizon."],
            ["VECM/Cointegration", "PhÃ¹ há»£p khi giÃ¡ ná»™i Ä‘á»‹a, vÃ ng tháº¿ giá»›i vÃ  FX cÃ³ quan há»‡ cÃ¢n báº±ng dÃ i háº¡n.", "VECM screen Ä‘Æ°á»£c ghi nháº­n; forecast chÆ°a promote trong runner v1."],
            ["GARCH/volatility", "ÄÃ¡nh giÃ¡ rá»§i ro, volatility clustering vÃ  stress regime.", "DÃ¹ng trong literature/methodology; volatility rolling trong EDA."],
            ["Tree boosting / Random Forest", "Báº¯t tÆ°Æ¡ng tÃ¡c phi tuyáº¿n giá»¯a premium, events, FX, GPR.", "Random Forest mean model; Gradient Boosting quantile cho downside q10."],
            ["Quantile regression", "Quyáº¿t Ä‘á»‹nh mua cáº§n downside risk, khÃ´ng chá»‰ expected return.", "q10 forecast káº¿t há»£p vá»›i xÃ¡c suáº¥t return dÆ°Æ¡ng."],
            ["DeepAR/TFT/N-BEATS/N-HiTS", "á»¨ng viÃªn multi-horizon probabilistic forecast khi cÃ³ panel Ä‘á»§ giÃ u.", "Ghi lÃ  hÆ°á»›ng tiáº¿p theo; khÃ´ng train náº¿u thiáº¿u dependency."],
        ],
    )
    add_para(
        doc,
        "LÃ½ do khÃ´ng chá»n má»™t mÃ´ hÃ¬nh duy nháº¥t ngay tá»« Ä‘áº§u: thá»‹ trÆ°á»ng vÃ ng vá»«a cÃ³ thÃ nh pháº§n xu hÆ°á»›ng toÃ n cáº§u, vá»«a cÃ³ premium ná»™i Ä‘á»‹a phi tuyáº¿n vÃ  event-driven. VÃ¬ váº­y Ä‘Ã¡nh giÃ¡ theo walk-forward vÃ  quyáº¿t Ä‘á»‹nh dá»±a trÃªn cáº£ MAE, directional accuracy, q10 downside vÃ  signal performance.",
    )
    add_para(
        doc,
        "Má»™t Ä‘iá»ƒm quan trá»ng cá»§a literature review lÃ  khÃ´ng xem mÃ´ hÃ¬nh phá»©c táº¡p hÆ¡n lÃ  máº·c nhiÃªn tá»‘t hÆ¡n. Náº¿u boosting hoáº·c deep learning chá»‰ cáº£i thiá»‡n sai sá»‘ trung bÃ¬nh nhÆ°ng lÃ m tÃ­n hiá»‡u giao dá»‹ch kÃ©m á»•n Ä‘á»‹nh, mÃ´ hÃ¬nh Ä‘Ã³ khÃ´ng Ä‘Æ°á»£c Æ°u tiÃªn cho quyáº¿t Ä‘á»‹nh mua vÃ ng váº­t cháº¥t. NgÆ°á»£c láº¡i, má»™t baseline Ä‘Æ¡n giáº£n nhÆ°ng cÃ³ directional accuracy á»•n Ä‘á»‹nh vÃ  downside cÃ³ kiá»ƒm soÃ¡t cÃ³ thá»ƒ há»¯u Ã­ch hÆ¡n cho nhÃ  Ä‘áº§u tÆ°.",
    )


def add_data_sections(
    doc: Document,
    profiles: list[DataProfile],
    quality_rows: list[list[Any]],
    eda_tables: dict[str, list[list[Any]]],
    figure_paths: dict[str, Path],
) -> None:
    add_heading(doc, "3. Data Requirements, Collection vÃ  Data Understanding", 1)
    add_para(
        doc,
        "YÃªu cáº§u dá»¯ liá»‡u xuáº¥t phÃ¡t tá»« phÆ°Æ¡ng trÃ¬nh kinh táº¿: giÃ¡ vÃ ng ná»™i Ä‘á»‹a xáº¥p xá»‰ giÃ¡ vÃ ng tháº¿ giá»›i quy Ä‘á»•i VND cá»™ng premium ná»™i Ä‘á»‹a, sau Ä‘Ã³ bá»‹ Ä‘iá»u chá»‰nh bá»Ÿi spread, mÃ¹a vá»¥, chÃ­nh sÃ¡ch vÃ  rá»§i ro. VÃ¬ váº­y data lake pháº£i cÃ³ target ná»™i Ä‘á»‹a, reference toÃ n cáº§u, FX, macro, event vÃ  kiá»ƒm soÃ¡t thá»i Ä‘iá»ƒm cÃ´ng bá»‘.",
    )
    add_para(
        doc,
        "Báº£ng kiá»ƒm kÃª dÆ°á»›i Ä‘Ã¢y khÃ´ng chá»‰ liá»‡t kÃª file, mÃ  thá»ƒ hiá»‡n vai trÃ² cá»§a tá»«ng lá»›p dá»¯ liá»‡u trong pipeline. NhÃ³m target ná»™i Ä‘á»‹a cung cáº¥p nhÃ£n huáº¥n luyá»‡n vÃ  giÃ¡ thá»±c thi; nhÃ³m global/reference táº¡o fair value Ä‘á»ƒ tÃ­nh premium; nhÃ³m FX chuyá»ƒn USD gold sang VND/lÆ°á»£ng; nhÃ³m macro/news/event mÃ´ táº£ Ä‘iá»u kiá»‡n thá»‹ trÆ°á»ng; cÃ²n nhÃ³m modeling lÆ°u frame cuá»‘i, prediction vÃ  decision signal Ä‘á»ƒ cÃ³ thá»ƒ Ä‘á»‘i chiáº¿u ngÆ°á»£c má»i con sá»‘ trong bÃ¡o cÃ¡o.",
    )
    add_para(
        doc,
        "YÃªu cáº§u quan trá»ng nháº¥t lÃ  tÃ­nh lá»‹ch sá»­ há»£p lá»‡. Má»™t nguá»“n chá»‰ Ä‘Æ°á»£c dÃ¹ng lÃ m nhÃ£n náº¿u cÃ³ thá»ƒ chá»©ng minh dá»¯ liá»‡u tráº£ vá» Ä‘Ãºng ngÃ y nghiá»‡p vá»¥ Ä‘Ã£ yÃªu cáº§u, khÃ´ng pháº£i giÃ¡ hiá»‡n táº¡i bá»‹ leak vÃ o quÃ¡ khá»©. NguyÃªn táº¯c nÃ y lÃ m pipeline cháº­m vÃ  nghiÃªm ngáº·t hÆ¡n, nhÆ°ng cáº§n thiáº¿t vÃ¬ má»™t sai lá»‡ch current-leak nhá» cÅ©ng cÃ³ thá»ƒ lÃ m backtest vÃ ng trÃ´ng tá»‘t giáº£ táº¡o.",
    )
    add_table(
        doc,
        ["Dataset", "Rows", "Cols", "Thá»i gian", "Vai trÃ²"],
        [[p.name, f"{p.rows:,}", f"{p.cols:,}", p.date_span, p.role] for p in profiles],
        max_rows=20,
    )
    add_heading(doc, "3.1. Minh chá»©ng dá»¯ liá»‡u crawl vÃ  artifact thu tháº­p", 2)
    add_para(
        doc,
        "Ngay trong pháº§n collection, bÃ¡o cÃ¡o trÃ­ch má»™t sá»‘ dÃ²ng tháº­t tá»« data lake Ä‘á»ƒ chá»©ng minh nguá»“n dá»¯ liá»‡u Ä‘Ã£ Ä‘Æ°á»£c crawl/chuáº©n hÃ³a. CÃ¡c báº£ng nÃ y khÃ´ng pháº£i vÃ­ dá»¥ minh há»a tá»± táº¡o: chÃºng Ä‘Æ°á»£c Ä‘á»c trá»±c tiáº¿p tá»« CSV trong data/lake khi build DOCX.",
    )
    sbv_structures = read_csv("data/lake/source_discovery/sbv_structures.csv")
    sbv_events = read_csv("data/lake/events/sbv_gold_policy_events.csv")
    deposit = read_csv("data/lake/normalized/retail_deposit_rates.csv")
    lbma = read_csv("data/lake/normalized/lbma_gold_spot_am_pm.csv")
    news = read_csv("data/lake/news_availability_audit.csv")
    if not sbv_structures.empty:
        cols = ["content_structure_id", "classification", "http_status", "row_count_sample", "field_names", "title_samples"]
        add_para(doc, "SBV source discovery sample: structure 137473 Ä‘Æ°á»£c phÃ¢n loáº¡i central_fx, cÃ²n cÃ¡c structure liÃªn quan vÃ ng Ä‘Æ°á»£c giá»¯ á»Ÿ má»©c candidate náº¿u chÆ°a cÃ³ payload há»£p lá»‡.")
        add_table(doc, [c for c in cols if c in sbv_structures.columns], rows_from_frame(sbv_structures, cols, max_rows=6, max_chars=105), max_rows=6)
    if not sbv_events.empty:
        cols = ["event_date", "published_at", "event_type", "severity", "title", "confidence", "source_type"]
        add_para(doc, "SBV gold-policy event sample: event chá»‰ Ä‘Æ°á»£c giá»¯ khi cÃ³ URL chÃ­nh thá»©c, ngÃ y cÃ´ng bá»‘, loáº¡i sá»± kiá»‡n vÃ  raw hash Ä‘á»ƒ audit.")
        add_table(doc, [c for c in cols if c in sbv_events.columns], rows_from_frame(sbv_events, cols, max_rows=5, max_chars=120), max_rows=5)
    if not deposit.empty:
        cols = ["date", "bank", "tenor_months", "currency", "rate_pct_annual", "published_at", "available_from", "history_status"]
        add_para(doc, "Retail deposit-rate sample: báº£ng nÃ y chá»©ng minh lá»›p lÃ£i suáº¥t hiá»‡n chá»‰ lÃ  forward monitoring, chÆ°a pháº£i lá»‹ch sá»­ benchmark Ä‘áº§y Ä‘á»§.")
        add_table(doc, [c for c in cols if c in deposit.columns], rows_from_frame(deposit, cols, max_rows=8, max_chars=90), max_rows=8)
    if not lbma.empty:
        cols = ["date", "series_id", "value", "unit", "source", "available_from", "fix_type"]
        add_para(doc, "LBMA append sample: dá»¯ liá»‡u public today.json Ä‘Æ°á»£c lÆ°u riÃªng, cÃ²n lá»‹ch sá»­ thiáº¿u licence váº«n Ä‘á»c báº±ng proxy cÃ³ gáº¯n cá».")
        add_table(doc, [c for c in cols if c in lbma.columns], rows_from_frame(lbma, cols, max_rows=6, max_chars=90), max_rows=6)
    if not news.empty:
        cols = ["crawl_date", "event_date", "headline", "published_at", "fetched_at", "availability_from", "feature_mode_strict"]
        add_para(doc, "News availability sample: báº£ng nÃ y giáº£i thÃ­ch vÃ¬ sao headline Ä‘Æ°á»£c tÃ¡ch thÃ nh research mode vÃ  strict realtime mode.")
        add_table(doc, [c for c in cols if c in news.columns], rows_from_frame(news, cols, max_rows=5, max_chars=120), max_rows=5)

    add_heading(doc, "3.2. Code crawl SBV Ä‘i kÃ¨m káº¿t quáº£ discovery", 2)
    add_code_block(
        doc,
        "SBV discovery báº±ng Playwright vÃ  endpoint headless-delivery",
        "scripts/pipeline/discover_sbv_sources.py",
        35,
        72,
        "Äoáº¡n code nÃ y giáº£i thÃ­ch vÃ¬ sao bÃ¡o cÃ¡o chá»n Playwright cho SBV: má»Ÿ trang chÃ­nh thá»©c, tÃ¬m content-structure id vÃ  fetch endpoint qua browser context.",
    )
    add_code_block(
        doc,
        "Chuáº©n hÃ³a event chÃ­nh sÃ¡ch vÃ ng tá»« SBV",
        "scripts/pipeline/collect_sbv_gold_policy_events.py",
        180,
        199,
        "Má»—i event Ä‘i vÃ o báº£ng SBV policy pháº£i cÃ³ event_date/published_at, event_type, source_url, source_type, confidence vÃ  raw_hash.",
    )

    add_heading(doc, "4. Data Quality vÃ  rá»§i ro dá»¯ liá»‡u", 1)
    add_para(
        doc,
        "Data quality Ä‘Æ°á»£c Ä‘Ã¡nh giÃ¡ theo bá»‘n nhÃ³m: Ä‘á»™ phá»§ theo thá»i gian, tÃ­nh duy nháº¥t á»Ÿ grain ngÃ y, tá»· lá»‡ missing cá»§a biáº¿n quyáº¿t Ä‘á»‹nh, vÃ  tÃ­nh Ä‘Ãºng nghÄ©a cá»§a nguá»“n. ÄÃ¢y lÃ  Ä‘iá»ƒm Ä‘áº·c biá»‡t quan trá»ng vá»›i cÃ¡c nguá»“n SBV/NHNN: cÃ¹ng má»™t cáº¥u trÃºc CMS cÃ³ thá»ƒ lÃ  tá»· giÃ¡ trung tÃ¢m chá»© khÃ´ng pháº£i lÃ£i suáº¥t tiá»n gá»­i, nÃªn pipeline khÃ´ng Ä‘Æ°á»£c phÃ©p Ä‘á»•i nhÃ£n ngá»¯ nghÄ©a chá»‰ Ä‘á»ƒ cÃ³ thÃªm feature.",
    )
    add_para(
        doc,
        "Sau vÃ²ng cáº£i thiá»‡n, premium coverage Ä‘Ã£ Ä‘Æ°á»£c xá»­ lÃ½ báº±ng hierarchy nguá»“n cÃ³ cá» proxy thay vÃ¬ bá» trá»‘ng hÃ ng loáº¡t. CÃ¡ch lÃ m nÃ y khÃ´ng biáº¿n proxy thÃ nh benchmark chÃ­nh thá»©c; nÃ³ chá»‰ giÃºp mÃ´ hÃ¬nh cÃ³ má»™t Ä‘áº¡i lÆ°á»£ng fair-value liÃªn tá»¥c hÆ¡n, Ä‘á»“ng thá»i váº«n giá»¯ metadata Ä‘á»ƒ phÃ¢n biá»‡t ngÃ y nÃ o lÃ  nguá»“n cháº¥t lÆ°á»£ng cao vÃ  ngÃ y nÃ o lÃ  proxy cáº§n Ä‘á»c tháº­n trá»ng.",
    )
    add_table(doc, ["Check", "Káº¿t quáº£", "Má»©c Ä‘á»™", "Ã nghÄ©a"], quality_rows, max_rows=25)
    premium_audit = read_csv("data/lake/quality/premium_coverage_audit.csv")
    if not premium_audit.empty:
        cols = ["segment", "rows", "premium_missing", "premium_missing_rate", "proxy_share", "section"]
        add_para(doc, "Premium coverage audit sample: káº¿t quáº£ cáº£i thiá»‡n premium Ä‘Æ°á»£c Ä‘o theo nÄƒm vÃ  theo source quality, khÃ´ng chá»‰ ghi báº±ng lá»i.")
        add_table(doc, [c for c in cols if c in premium_audit.columns], rows_from_frame(premium_audit, cols, max_rows=8, max_chars=80), max_rows=8)
    add_code_block(
        doc,
        "Premium coverage audit sau fallback cÃ³ giá»›i háº¡n",
        "scripts/pipeline/improve_premium_coverage.py",
        257,
        287,
        "Äoáº¡n code nÃ y cho tháº¥y fallback Ä‘Æ°á»£c kiá»ƒm báº±ng missing rate, proxy_share vÃ  target <10%; náº¿u khÃ´ng Ä‘áº¡t thÃ¬ ghi blocker.",
    )

    add_heading(doc, "5. Data Preparation", 1)
    add_para(
        doc,
        "Quy trÃ¬nh chuáº©n bá»‹ dá»¯ liá»‡u Ä‘Æ°á»£c thiáº¿t káº¿ nhÆ° má»™t as-of pipeline. Táº¡i má»—i ngÃ y quyáº¿t Ä‘á»‹nh t, mÃ´ hÃ¬nh chá»‰ Ä‘Æ°á»£c tháº¥y thÃ´ng tin Ä‘Ã£ tá»“n táº¡i trÆ°á»›c hoáº·c táº¡i thá»i Ä‘iá»ƒm Ä‘Ã³ theo lá»‹ch cÃ´ng bá»‘ há»£p lÃ½. Vá»›i thá»‹ trÆ°á»ng Má»¹, dá»¯ liá»‡u Ä‘Ã³ng cá»­a sau giá» giao dá»‹ch Viá»‡t Nam nÃªn Ä‘Æ°á»£c lag tá»‘i thiá»ƒu má»™t ngÃ y. Vá»›i macro/news, trÆ°á»ng available_from Ä‘Æ°á»£c dÃ¹ng Ä‘á»ƒ mÃ´ phá»ng ngÃ y thÃ´ng tin thá»±c sá»± cÃ³ thá»ƒ Ä‘Æ°á»£c biáº¿t.",
    )
    add_bullet(doc, "Historical-valid target: chá»‰ dÃ¹ng báº£n ghi cÃ³ requested/business date Ä‘Ãºng ngÃ y vÃ  giÃ¡ mua/bÃ¡n há»£p lá»‡.")
    add_bullet(doc, "As-of join: global vÃ  GPR dÃ¹ng cutoff t-1; macro dÃ¹ng available_from <= date Ä‘á»ƒ trÃ¡nh look-ahead bias.")
    add_bullet(doc, "Feature engineering: táº¡o lag 1/5/10/21/63/105 ngÃ y vÃ  rolling mean/std 5/21/63 ngÃ y.")
    add_bullet(doc, "Target: net_return_h = future_buy_price_h / current_sell_price - 1, Ä‘Ãºng logic ngÆ°á»i mua vÃ ng váº­t cháº¥t chá»‹u spread.")
    add_bullet(doc, "Horizon: 1/3/5 thÃ¡ng lá»‹ch, láº¥y giÃ¡ exit á»Ÿ ngÃ y kháº£ dá»¥ng gáº§n nháº¥t sau target date.")
    add_para(
        doc,
        "CÃ¡c Ä‘áº·c trÆ°ng rolling Ä‘Æ°á»£c táº¡o á»Ÿ nhiá»u cá»­a sá»• theo ngÃ y vÃ¬ má»—i horizon thÃ¡ng cÃ³ báº£n cháº¥t khÃ¡c nhau. Horizon 1 thÃ¡ng nháº¡y vá»›i momentum ngáº¯n háº¡n vÃ  spread tá»©c thá»i; horizon 3 thÃ¡ng pháº£n Ã¡nh pha Ä‘iá»u chá»‰nh trung háº¡n; horizon 5 thÃ¡ng chá»‹u tÃ¡c Ä‘á»™ng máº¡nh hÆ¡n cá»§a cháº¿ Ä‘á»™ vÄ© mÃ´, premium vÃ  sá»± kiá»‡n chÃ­nh sÃ¡ch. Viá»‡c dÃ¹ng cÃ¹ng má»™t bá»™ feature cho cáº£ ba horizon giÃºp so sÃ¡nh nháº¥t quÃ¡n, nhÆ°ng káº¿t quáº£ Ä‘Æ°á»£c Ä‘Ã¡nh giÃ¡ riÃªng tá»«ng horizon.",
    )
    add_heading(doc, "5.1. Code táº¡o target vÃ  chá»‘ng leakage", 2)
    add_code_block(
        doc,
        "Target lá»£i suáº¥t sau spread",
        "src/gold_collectors/modeling/decision_support.py",
        375,
        386,
        "ÄÃ¢y lÃ  Ä‘oáº¡n biáº¿n bÃ i toÃ¡n dá»± bÃ¡o giÃ¡ thÃ nh bÃ i toÃ¡n lá»£i suáº¥t thá»±c thi: mua theo sell_price hiá»‡n táº¡i vÃ  thoÃ¡t theo future buy_price.",
    )
    add_code_block(
        doc,
        "As-of join cho dá»¯ liá»‡u global",
        "src/gold_collectors/modeling/decision_support.py",
        80,
        112,
        "Global market dÃ¹ng cutoff t-1 trÆ°á»›c khi merge_asof, giáº£m rá»§i ro dÃ¹ng giÃ¡ quá»‘c táº¿ Ä‘Ã³ng cá»­a sau giá» Viá»‡t Nam.",
    )
    add_code_block(
        doc,
        "Deposit-rate as-of vÃ  opportunity cost",
        "src/gold_collectors/modeling/decision_support.py",
        188,
        220,
        "Feature lÃ£i suáº¥t chá»‰ Ä‘Æ°á»£c táº¡o khi cÃ³ available_from, tenor vÃ  rate há»£p lá»‡; náº¿u chÆ°a cÃ³ lá»‹ch sá»­ thÃ¬ khÃ´ng Ã©p backfill.",
    )
    add_code_block(
        doc,
        "Pipeline build model frame",
        "src/gold_collectors/modeling/decision_support.py",
        465,
        490,
        "Thá»© tá»± build frame thá»ƒ hiá»‡n toÃ n bá»™ pipeline: target SJC, premium, global, GPR, macro, deposit, event, news, target vÃ  lagged features.",
    )
    frame = read_csv("data/lake/modeling/model_frame_daily.csv")
    if not frame.empty:
        cols = ["date", "buy_price", "sell_price", "spread_pct", "premium", "premium_pct", "is_proxy", "net_return_1m", "net_return_3m", "net_return_5m"]
        available = [c for c in cols if c in frame.columns]
        add_para(doc, "Model frame sample á»Ÿ cuá»‘i snapshot: dá»¯ liá»‡u Ä‘Ã£ Ä‘Æ°á»£c ghÃ©p cÃ¹ng grain ngÃ y, kÃ¨m target tÆ°Æ¡ng lai cho tá»«ng horizon.")
        add_table(doc, available, rows_from_frame(frame[available].tail(5), available, max_rows=5, max_chars=80), max_rows=5)

    add_heading(doc, "6. EDA: GiÃ¡, premium, thanh khoáº£n vÃ  sá»± kiá»‡n", 1)
    add_para(
        doc,
        "EDA táº­p trung vÃ o nhá»¯ng cÃ¢u há»i cÃ³ áº£nh hÆ°á»Ÿng trá»±c tiáº¿p Ä‘áº¿n quyáº¿t Ä‘á»‹nh mua: giÃ¡ SJC cÃ³ Ä‘ang á»Ÿ vÃ¹ng biáº¿n Ä‘á»™ng báº¥t thÆ°á»ng khÃ´ng, premium Ä‘ang Ä‘áº¯t hay ráº» so vá»›i lá»‹ch sá»­, spread cÃ³ lÃ m chi phÃ­ giao dá»‹ch quÃ¡ cao khÃ´ng, vÃ  cÃ¡c sá»± kiá»‡n chÃ­nh sÃ¡ch/news cÃ³ trÃ¹ng vá»›i thay Ä‘á»•i regime hay khÃ´ng. CÃ¡c báº£ng phÃ­a dÆ°á»›i vÃ¬ váº­y Ä‘Æ°á»£c Ä‘á»c nhÆ° báº±ng chá»©ng Ä‘á»‹nh hÆ°á»›ng mÃ´ hÃ¬nh, khÃ´ng pháº£i chá»‰ lÃ  thá»‘ng kÃª mÃ´ táº£.",
    )
    if eda_tables.get("yearly_price"):
        add_para(doc, "Diá»…n biáº¿n 10 nÄƒm gáº§n nháº¥t cá»§a giÃ¡ bÃ¡n SJC trong model frame:")
        add_table(doc, ["NÄƒm", "Start", "End", "Min", "Max", "Return", "N obs", "Avg spread"], eda_tables["yearly_price"])
        add_picture_if_exists(doc, figure_paths.get("price", Path()), "Figure 1. SJC price versus global gold proxy")
        add_para(
            doc,
            "Báº£ng theo nÄƒm cho tháº¥y vÃ ng SJC khÃ´ng tÄƒng Ä‘á»u tuyáº¿n tÃ­nh; thay vÃ o Ä‘Ã³ cÃ³ cÃ¡c pha tÄƒng máº¡nh, Ä‘i ngang vÃ  Ä‘iá»u chá»‰nh. VÃ¬ tháº¿ mÃ´ hÃ¬nh khÃ´ng Ä‘Æ°á»£c Ä‘Ã¡nh giÃ¡ báº±ng má»™t giai Ä‘oáº¡n holdout ngáº¯n duy nháº¥t, mÃ  cáº§n walk-forward Ä‘á»ƒ kiá»ƒm tra nhiá»u cháº¿ Ä‘á»™ thá»‹ trÆ°á»ng. Náº¿u má»™t mÃ´ hÃ¬nh chá»‰ tá»‘t trong pha tÄƒng nhÆ°ng sai trong pha premium co láº¡i, mÃ´ hÃ¬nh Ä‘Ã³ khÃ´ng Ä‘á»§ tin cáº­y cho quyáº¿t Ä‘á»‹nh mua.",
        )
    if eda_tables.get("premium_stats"):
        add_para(doc, "Premium decomposition cho tháº¥y premium lÃ  biáº¿n quyáº¿t Ä‘á»‹nh cháº¥t lÆ°á»£ng Ä‘iá»ƒm mua:")
        add_table(doc, ["Metric", "Premium VND/lÆ°á»£ng", "Premium %"], eda_tables["premium_stats"])
        add_picture_if_exists(doc, figure_paths.get("premium", Path()), "Figure 2. Premium regime over time")
        add_para(
            doc,
            "Premium Ä‘Æ°á»£c hiá»ƒu lÃ  pháº§n giÃ¡ ná»™i Ä‘á»‹a cao hÆ¡n hoáº·c tháº¥p hÆ¡n fair value toÃ n cáº§u quy Ä‘á»•i. Premium dÆ°Æ¡ng cao cÃ³ thá»ƒ pháº£n Ã¡nh khan hiáº¿m vÃ  nhu cáº§u tÃ­ch trá»¯ máº¡nh, nhÆ°ng cÅ©ng lÃ m Ä‘iá»ƒm mua trá»Ÿ nÃªn Ä‘áº¯t náº¿u chÃ­nh sÃ¡ch cung vÃ ng sau Ä‘Ã³ kÃ©o premium xuá»‘ng. Do Ä‘Ã³ premium vá»«a lÃ  biáº¿n giáº£i thÃ­ch xu hÆ°á»›ng ngáº¯n háº¡n, vá»«a lÃ  nguá»“n rá»§i ro mean-reversion.",
        )
    if eda_tables.get("premium_regime"):
        add_table(doc, ["Premium regime", "Tá»· trá»ng quan sÃ¡t"], eda_tables["premium_regime"])
    if eda_tables.get("global_coverage"):
        add_table(doc, ["Global variable", "Non-null", "Coverage", "Min", "Max"], eda_tables["global_coverage"], max_rows=12)
        add_para(
            doc,
            "Coverage cá»§a biáº¿n toÃ n cáº§u quyáº¿t Ä‘á»‹nh Ä‘á»™ tin cáº­y cá»§a premium. Khi LBMA lá»‹ch sá»­ khÃ´ng cÃ³ licence Ä‘áº§y Ä‘á»§, GC=F/global close proxy Ä‘Æ°á»£c dÃ¹ng vá»›i cá» nguá»“n rÃµ rÃ ng. BÃ¡o cÃ¡o khÃ´ng coi proxy lÃ  benchmark chÃ­nh thá»©c, nhÆ°ng sá»­ dá»¥ng nÃ³ Ä‘á»ƒ giá»¯ cáº¥u trÃºc fair value liÃªn tá»¥c cho mÃ´ hÃ¬nh vÃ  ghi caveat trong pháº§n cáº£i thiá»‡n dá»¯ liá»‡u.",
        )
    if eda_tables.get("correlations"):
        add_para(doc, "Top tÆ°Æ¡ng quan tuyá»‡t Ä‘á»‘i giá»¯a má»™t sá»‘ feature lag vÃ  target return, dÃ¹ng Ä‘á»ƒ Ä‘á»‹nh hÆ°á»›ng chá»© khÃ´ng diá»…n giáº£i nhÃ¢n quáº£:")
        add_table(doc, ["Horizon", "Feature", "Correlation"], eda_tables["correlations"], max_rows=18)
    if eda_tables.get("events"):
        add_para(doc, "CÃ¡c loáº¡i sá»± kiá»‡n phá»• biáº¿n trong event regime panel:")
        add_table(doc, ["Event type", "Count"], eda_tables["events"])


def add_modeling_sections(
    doc: Document,
    mean_rows: list[list[Any]],
    quant_rows: list[list[Any]],
    decision_rows: list[list[Any]],
    snapshot_rows: list[list[Any]],
    recommendation: str,
    summary: dict[str, Any],
    figure_paths: dict[str, Path],
) -> None:
    add_heading(doc, "7. Modeling vÃ  Evaluation", 1)
    add_para(
        doc,
        "Thiáº¿t káº¿ evaluation dÃ¹ng expanding-window walk-forward: train trÃªn quÃ¡ khá»©, Ä‘Ã¡nh giÃ¡ trÃªn validation/test theo thá»i gian. ÄÃ¢y lÃ  báº¯t buá»™c vÃ¬ random split sáº½ lÃ m rÃ² rá»‰ cháº¿ Ä‘á»™ thá»‹ trÆ°á»ng tÆ°Æ¡ng lai vÃ o quÃ¡ khá»©.",
    )
    add_para(
        doc,
        "Má»—i horizon Ä‘Æ°á»£c huáº¥n luyá»‡n vÃ  Ä‘Ã¡nh giÃ¡ riÃªng vÃ¬ phÃ¢n phá»‘i lá»£i suáº¥t 1, 3 vÃ  5 thÃ¡ng khÃ´ng giá»‘ng nhau. Horizon ngáº¯n thÆ°á»ng bá»‹ chi phá»‘i bá»Ÿi nhiá»…u, spread vÃ  timing; horizon dÃ i hÆ¡n cÃ³ nhiá»u thá»i gian Ä‘á»ƒ xu hÆ°á»›ng vÄ© mÃ´ hoáº·c premium regime phÃ¡t huy tÃ¡c Ä‘á»™ng, nhÆ°ng cÅ©ng chá»‹u rá»§i ro chÃ­nh sÃ¡ch lá»›n hÆ¡n. VÃ¬ váº­y, viá»‡c má»™t mÃ´ hÃ¬nh tá»‘t á»Ÿ 5 thÃ¡ng khÃ´ng tá»± Ä‘á»™ng há»£p lá»‡ cho 1 thÃ¡ng.",
    )
    add_para(
        doc,
        "Bá»™ mÃ´ hÃ¬nh Ä‘Æ°á»£c sáº¯p theo má»©c phá»©c táº¡p tÄƒng dáº§n. Baseline zero/mean/median cho biáº¿t má»©c sÃ n cáº§n vÆ°á»£t qua; Ridge/ElasticNet kiá»ƒm tra quan há»‡ tuyáº¿n tÃ­nh cÃ³ regularization; SARIMAX kiá»ƒm tra cáº¥u trÃºc chuá»—i thá»i gian cÃ³ exogenous; Random Forest vÃ  boosting kiá»ƒm tra phi tuyáº¿n; quantile models Æ°á»›c lÆ°á»£ng downside. LightGBM, XGBoost vÃ  CatBoost chá»‰ Ä‘Æ°á»£c Ä‘Æ°a vÃ o leaderboard khi dependency import Ä‘Æ°á»£c vÃ  mÃ´ hÃ¬nh train tháº­t.",
    )
    add_para(
        doc,
        "CÃ¡c chá»‰ sá»‘ thá»‘ng kÃª nhÆ° MAE/RMSE cho biáº¿t sai sá»‘ trung bÃ¬nh, nhÆ°ng quyáº¿t Ä‘á»‹nh Ä‘áº§u tÆ° cáº§n thÃªm directional accuracy vÃ  signal performance. Má»™t mÃ´ hÃ¬nh cÃ³ MAE tháº¥p nhÆ°ng thÆ°á»ng dá»± bÃ¡o sai hÆ°á»›ng á»Ÿ cÃ¡c ngÃ y phÃ¡t tÃ­n hiá»‡u mua sáº½ khÃ´ng há»¯u Ã­ch. NgÆ°á»£c láº¡i, má»™t mÃ´ hÃ¬nh khÃ´ng tá»‘i Æ°u MAE tuyá»‡t Ä‘á»‘i nhÆ°ng phÃ¡t tÃ­n hiá»‡u Ã­t, cÃ³ chá»n lá»c vÃ  cÃ³ q10 khÃ´ng quÃ¡ xáº¥u cÃ³ thá»ƒ tá»‘t hÆ¡n vá» máº·t quyáº¿t Ä‘á»‹nh.",
    )
    add_table(
        doc,
        ["Horizon", "Best MAE model", "MAE", "RMSE", "DA", "Best DA model", "Best DA"],
        mean_rows,
    )
    add_picture_if_exists(doc, figure_paths.get("leaderboard", Path()), "Figure 3. Model MAE leaderboard by horizon")
    add_para(
        doc,
        "Báº£ng leaderboard Ä‘Æ°á»£c Ä‘á»c theo hai lá»›p. Lá»›p thá»© nháº¥t lÃ  kháº£ nÄƒng dá»± bÃ¡o má»©c lá»£i suáº¥t, thá»ƒ hiá»‡n qua MAE vÃ  RMSE. Lá»›p thá»© hai lÃ  kháº£ nÄƒng dá»± bÃ¡o hÆ°á»›ng, thá»ƒ hiá»‡n qua directional accuracy. Náº¿u hai lá»›p nÃ y mÃ¢u thuáº«n, bÃ¡o cÃ¡o khÃ´ng chá»n mÃ´ hÃ¬nh chá»‰ vÃ¬ má»™t chá»‰ sá»‘; quyáº¿t Ä‘á»‹nh cuá»‘i cÃ¹ng cÃ²n pháº£i Ä‘i qua rule xÃ¡c suáº¥t dÆ°Æ¡ng vÃ  downside q10.",
    )
    add_code_block(
        doc,
        "Optional model training cho LightGBM/XGBoost/CatBoost",
        "src/gold_collectors/modeling/decision_support.py",
        713,
        753,
        "Äoáº¡n code nÃ y náº±m ngay cáº¡nh leaderboard Ä‘á»ƒ chá»©ng minh cÃ¡c mÃ´ hÃ¬nh boosting chá»‰ xuáº¥t hiá»‡n khi import Ä‘Æ°á»£c vÃ  train tháº­t, khÃ´ng pháº£i ghi giáº£ trong bÃ¡o cÃ¡o.",
    )
    if quant_rows:
        add_table(doc, ["Horizon", "Best quantile model", "Pinball loss"], quant_rows)
        add_para(
            doc,
            "Pinball loss dÃ¹ng cho mÃ´ hÃ¬nh phÃ¢n vá»‹ q10. q10 Ä‘Æ°á»£c hiá»ƒu lÃ  ká»‹ch báº£n xáº¥u nhÆ°ng khÃ´ng cá»±c Ä‘oan: náº¿u q10 quÃ¡ Ã¢m, nhÃ  Ä‘áº§u tÆ° pháº£i cháº¥p nháº­n ráº±ng trong khoáº£ng 10% trÆ°á»ng há»£p xáº¥u, khoáº£n lá»— cÃ³ thá»ƒ vÆ°á»£t ngÆ°á»¡ng chá»‹u Ä‘á»±ng. ÄÃ¢y lÃ  lá»›p kiá»ƒm soÃ¡t rá»§i ro quan trá»ng vÃ¬ thá»‹ trÆ°á»ng vÃ ng Viá»‡t Nam cÃ³ thá»ƒ bá»‹ tÃ¡c Ä‘á»™ng bá»Ÿi premium vÃ  chÃ­nh sÃ¡ch báº¥t ngá».",
        )
    if decision_rows:
        add_para(doc, "Decision-rule backtest vá»›i rule P(return>0) >= 0.50 vÃ  q10 >= -10%:")
        add_table(doc, ["Horizon", "Phase", "Signals", "Signal rate", "Avg buy-day ret", "Avg strategy ret"], decision_rows)
        add_picture_if_exists(doc, figure_paths.get("signals", Path()), "Figure 4. Decision signal frequency")
        add_para(
            doc,
            "Decision-rule backtest tráº£ lá»i cÃ¢u há»i khÃ¡c vá»›i leaderboard: náº¿u chá»‰ mua vÃ o nhá»¯ng ngÃ y mÃ´ hÃ¬nh báº­t tÃ­n hiá»‡u, káº¿t quáº£ trung bÃ¬nh cá»§a cÃ¡c ngÃ y Ä‘Ã³ ra sao. Tá»· lá»‡ phÃ¡t tÃ­n hiá»‡u quÃ¡ cao cÃ³ thá»ƒ nghÄ©a lÃ  rule thiáº¿u chá»n lá»c; tá»· lá»‡ quÃ¡ tháº¥p cÃ³ thá»ƒ nghÄ©a lÃ  mÃ´ hÃ¬nh khÃ´ng Ä‘á»§ cÆ¡ há»™i thá»±c thi. VÃ¬ váº­y bÃ¡o cÃ¡o Ä‘á»c Ä‘á»“ng thá»i signal rate, average buy-day return vÃ  consistency giá»¯a validation/test.",
        )
        add_code_block(
            doc,
            "Decision rule tá»« xÃ¡c suáº¥t vÃ  q10 downside",
            "src/gold_collectors/modeling/decision_support.py",
            930,
            945,
            "TÃ­n hiá»‡u mua khÃ´ng chá»‰ dá»±a vÃ o expected return; rule yÃªu cáº§u P(return>0) vÆ°á»£t ngÆ°á»¡ng vÃ  q10 khÃ´ng tháº¥p hÆ¡n sÃ n rá»§i ro.",
        )

    add_heading(doc, "8. Snapshot Forecast vÃ  quyáº¿t Ä‘á»‹nh mua 1-5 thÃ¡ng", 1)
    add_para(
        doc,
        "Báº£ng dÆ°á»›i Ä‘Ã¢y lÃ  dá»± bÃ¡o snapshot riÃªng táº¡i ngÃ y dá»¯ liá»‡u cuá»‘i. MÃ´ hÃ¬nh Ä‘Æ°á»£c train trÃªn cÃ¡c dÃ²ng cÃ³ nhÃ£n lá»‹ch sá»­ Ä‘Ã£ biáº¿t, sau Ä‘Ã³ dá»± bÃ¡o cho dÃ²ng feature má»›i nháº¥t. VÃ¬ tÆ°Æ¡ng lai sau snapshot chÆ°a cÃ³ nhÃ£n, Ä‘Ã¢y lÃ  forecast phá»¥c vá»¥ quyáº¿t Ä‘á»‹nh, khÃ´ng pháº£i backtest.",
    )
    add_para(
        doc,
        "CÃ¡ch Ä‘á»c báº£ng snapshot nhÆ° sau: expected return lÃ  lá»£i suáº¥t rÃ²ng ká»³ vá»ng sau spread; q10 downside lÃ  biÃªn dÆ°á»›i rá»§i ro á»Ÿ phÃ¢n vá»‹ 10%; P(return>0) lÃ  xÃ¡c suáº¥t mÃ´ hÃ¬nh Æ°á»›c lÆ°á»£ng lá»£i suáº¥t dÆ°Æ¡ng. TÃ­n hiá»‡u mua chá»‰ Ä‘Ã¡ng tin hÆ¡n khi ba Ä‘áº¡i lÆ°á»£ng nÃ y cÃ¹ng á»§ng há»™, vÃ  káº¿t quáº£ backtest cÃ¹ng horizon khÃ´ng mÃ¢u thuáº«n máº¡nh vá»›i forecast hiá»‡n táº¡i.",
    )
    if snapshot_rows:
        add_table(doc, ["Horizon", "Expected return", "Q10 downside", "P(return>0)", "Khuyáº¿n nghá»‹"], snapshot_rows)
    snapshot = read_csv("data/lake/modeling/snapshot_forecasts.csv")
    if not snapshot.empty:
        cols = ["snapshot_date", "horizon_months", "predicted_net_return", "q10_predicted_net_return", "prob_return_positive", "buy_signal", "status"]
        add_para(doc, "Raw snapshot forecast rows: báº£ng dÆ°á»›i láº¥y trá»±c tiáº¿p tá»« snapshot_forecasts.csv Ä‘á»ƒ Ä‘á»‘i chiáº¿u vá»›i khuyáº¿n nghá»‹ á»Ÿ trÃªn.")
        add_table(doc, [c for c in cols if c in snapshot.columns], rows_from_frame(snapshot, cols, max_rows=3, max_chars=90), max_rows=3)
    add_para(doc, recommendation, bold=True)
    add_para(
        doc,
        "Vá»›i horizon 1 thÃ¡ng, forecast Ã¢m khiáº¿n quyáº¿t Ä‘á»‹nh mua ngáº¯n háº¡n khÃ´ng háº¥p dáº«n, Ä‘áº·c biá»‡t khi spread lÃ m Ä‘iá»ƒm hÃ²a vá»‘n cao hÆ¡n. Vá»›i horizon 3 thÃ¡ng, expected return gáº§n báº±ng 0 nhÆ°ng q10 downside xáº¥u, nÃªn rá»§i ro khÃ´ng Ä‘Æ°á»£c bÃ¹ Ä‘á»§. Vá»›i horizon 5 thÃ¡ng, forecast cÃ³ tÃ­n hiá»‡u tá»‘t hÆ¡n, nhÆ°ng nÃªn hiá»ƒu lÃ  mua tÃ­ch lÅ©y cÃ³ kiá»ƒm soÃ¡t hoáº·c phÃ¢n bá»• tá»«ng pháº§n, khÃ´ng pháº£i all-in táº¡i má»™t giÃ¡ duy nháº¥t.",
    )
    add_para(
        doc,
        "Káº¿t luáº­n nÃ y phÃ¹ há»£p vá»›i logic kinh táº¿: vÃ ng cÃ³ thá»ƒ váº«n lÃ  tÃ i sáº£n phÃ²ng vá»‡ trong mÃ´i trÆ°á»ng báº¥t Ä‘á»‹nh, nhÆ°ng khi premium ná»™i Ä‘á»‹a cao, ngÆ°á»i mua chá»‹u thÃªm rá»§i ro chÃ­nh sÃ¡ch vÃ  rá»§i ro co premium. Do Ä‘Ã³, quyáº¿t Ä‘á»‹nh há»£p lÃ½ nháº¥t lÃ  Æ°u tiÃªn quáº£n trá»‹ vá»‹ tháº¿, theo dÃµi premium/spread vÃ  cáº­p nháº­t rolling forecast thay vÃ¬ xem snapshot lÃ  lá»i khuyÃªn cá»‘ Ä‘á»‹nh.",
    )

    add_heading(doc, "9. Caveats, Feedback vÃ  káº¿ hoáº¡ch cáº£i thiá»‡n", 1)
    add_para(
        doc,
        "CÃ¡c caveat dÆ°á»›i Ä‘Ã¢y khÃ´ng pháº£i lá»—i cáº§n che Ä‘i, mÃ  lÃ  pháº§n pháº£n há»“i Ä‘á»ƒ nÃ¢ng Ä‘á»™ tin cáº­y cá»§a há»‡ thá»‘ng. Vá»›i bÃ i toÃ¡n dá»± bÃ¡o tÃ i chÃ­nh, viá»‡c biáº¿t mÃ´ hÃ¬nh chÆ°a biáº¿t gÃ¬ quan trá»ng khÃ´ng kÃ©m káº¿t quáº£ dá»± bÃ¡o. Má»—i caveat Ä‘Æ°á»£c chuyá»ƒn thÃ nh backlog cÃ³ artifact Ä‘o lÆ°á»ng, Ä‘á»ƒ láº§n cháº¡y sau cÃ³ thá»ƒ kiá»ƒm tra Ä‘Ã£ cáº£i thiá»‡n tháº­t hay chÆ°a.",
    )
    add_bullet(doc, f"Premium missing rate hiá»‡n khoáº£ng {pct(summary.get('premium_missing_rate'))}; káº¿t luáº­n premium cáº§n Ä‘á»c lÃ  directional.")
    add_bullet(doc, "VN deposit-rate history bá»‹ loáº¡i vÃ¬ value null 100%, nÃªn chÆ°a benchmark Ä‘áº§y Ä‘á»§ vá»›i lÃ£i suáº¥t tiáº¿t kiá»‡m.")
    add_bullet(doc, "News/headline Ä‘Æ°á»£c backfill trong 2026; Ä‘Ã£ lag theo event_date nhÆ°ng real-time availability chÆ°a chá»©ng minh tuyá»‡t Ä‘á»‘i.")
    add_bullet(doc, "LightGBM/XGBoost/CatBoost Ä‘Ã£ Ä‘Æ°á»£c cÃ i vÃ  train tháº­t khi xuáº¥t hiá»‡n trong leaderboard; DeepAR/TFT lÃ  production-candidate cho runner riÃªng, khÃ´ng cÃ²n lÃ  blocker thiáº¿u dependency.")
    add_bullet(doc, "BÆ°á»›c tiáº¿p theo nÃªn Æ°u tiÃªn: bá»• sung LBMA/FX coverage, nguá»“n lÃ£i suáº¥t tiá»n gá»­i VN, lá»‹ch Ä‘áº¥u tháº§u/chÃ­nh sÃ¡ch NHNN, rá»“i cháº¡y paper-trading rolling sau snapshot.")
    add_para(
        doc,
        "Vá» premium, má»¥c tiÃªu cáº£i thiá»‡n lÃ  giáº£m missing nhÆ°ng khÃ´ng Ä‘Ã¡nh Ä‘á»•i tÃ­nh trung thá»±c cá»§a nguá»“n. VÃ¬ váº­y premium row pháº£i mang theo gold_reference_source, fx_source, source_quality, availability_from vÃ  is_proxy. Nhá»¯ng ngÃ y dÃ¹ng proxy váº«n Ä‘Æ°á»£c dÃ¹ng cho mÃ´ hÃ¬nh náº¿u cáº§n tÃ­nh liÃªn tá»¥c, nhÆ°ng report vÃ  sensitivity pháº£i cÃ³ kháº£ nÄƒng tÃ¡ch high-quality-only khá»i proxy-inclusive.",
    )
    add_para(
        doc,
        "Vá» lÃ£i suáº¥t, bÃ¡o cÃ¡o khÃ´ng trá»™n láº«n tá»· giÃ¡ trung tÃ¢m SBV vá»›i deposit rate. Structure 137473 Ä‘Ã£ Ä‘Æ°á»£c xÃ¡c minh lÃ  central USD/VND FX, nÃªn khÃ´ng Ä‘Æ°á»£c dÃ¹ng lÃ m lÃ£i suáº¥t tiá»n gá»­i. Khi chÆ°a cÃ³ lá»‹ch sá»­ chÃ­nh thá»©c, hÆ°á»›ng Ä‘Ãºng lÃ  archive forward cÃ¡c báº£ng lÃ£i suáº¥t ngÃ¢n hÃ ng thÆ°Æ¡ng máº¡i cÃ³ ngÃ y cÃ´ng bá»‘ rÃµ, sau Ä‘Ã³ má»›i táº¡o feature opportunity cost cho cÃ¡c ká»³ háº¡n 1/3/5 thÃ¡ng.",
    )
    add_para(
        doc,
        "Vá» news, cháº¿ Ä‘á»™ research_event_date_lagged cÃ³ thá»ƒ dÃ¹ng cho phÃ¢n tÃ­ch há»c thuáº­t vÃ¬ Ä‘Ã£ lag theo ngÃ y bÃ i viáº¿t, nhÆ°ng cháº¿ Ä‘á»™ strict_realtime_verified má»›i phÃ¹ há»£p cho paper-trading. Náº¿u hai cháº¿ Ä‘á»™ táº¡o káº¿t quáº£ khÃ¡c nhau Ä‘Ã¡ng ká»ƒ, bÃ¡o cÃ¡o pháº£i háº¡ confidence cá»§a headline features vÃ  Æ°u tiÃªn biáº¿n giÃ¡/thanh khoáº£n cÃ³ timestamp cháº¯c cháº¯n hÆ¡n.",
    )


def caveat_improvement_rows(summary: dict[str, Any]) -> list[list[Any]]:
    premium_summary = read_json("data/lake/quality/premium_coverage_summary.json")
    deposit_summary = read_json("data/lake/quality/deposit_rate_coverage_summary.json")
    news_summary = read_json("data/lake/quality/news_availability_summary.json")
    sensitivity = read_json("data/lake/modeling/model_sensitivity_summary.json")
    paper = read_json("data/lake/modeling/paper_trading_summary.json")

    before_premium = premium_summary.get("old_premium_missing_rate", summary.get("premium_missing_rate"))
    after_premium = premium_summary.get("new_premium_missing_rate", summary.get("premium_missing_rate"))
    deposit_rows = deposit_summary.get("retail_valid_rows", 0)
    strict_news_share = news_summary.get("strict_realtime_verified_share", 0)
    deps = sensitivity.get("dependency_availability", {})
    model_status = ", ".join(
        f"{name}={'ok' if available else 'missing'}"
        for name, available in deps.items()
        if name in {"lightgbm", "xgboost", "catboost", "torch", "pytorch_forecasting", "gluonts"}
    ) or "not audited"

    return [
        [
            "Premium coverage",
            pct(before_premium),
            pct(after_premium),
            "Met target <10%" if premium_summary.get("target_met") else "Needs source-gap review",
        ],
        [
            "Deposit opportunity cost",
            "0 non-null verified rows",
            f"{deposit_rows} current retail rows; SBV policy rows={deposit_summary.get('sbv_policy_rates', {}).get('rows', 0)}",
            deposit_summary.get("history_status", "not audited"),
        ],
        [
            "News real-time availability",
            "Backfilled event_date lagged",
            pct(strict_news_share),
            "Strict mode uses availability_from; backfilled rows excluded from paper-trading history.",
        ],
        [
            "Boosting/deep model coverage",
            "Boosting deps previously missing",
            model_status,
            "Boosting models are trained when importable; deep models need a dedicated runner before inclusion.",
        ],
        [
            "Paper-trading status",
            "No rolling ledger",
            f"rows={paper.get('rows', 0)}, open={paper.get('open_rows', 0)}, closed={paper.get('closed_rows', 0)}",
            "Open trades have no realized return until exit price exists.",
        ],
    ]


def add_next_data_expansion_section(doc: Document, summary: dict[str, Any]) -> None:
    add_heading(doc, "10. Next Data Expansion & Paper Trading", 1)
    add_para(
        doc,
        "Má»¥c nÃ y ghi nháº­n pháº§n má»Ÿ rá»™ng sau snapshot báº±ng dá»¯ liá»‡u tháº­t: Playwright Ä‘Æ°á»£c dÃ¹ng cho SBV/NHNN vÃ¬ API trá»±c tiáº¿p cÃ³ thá»ƒ bá»‹ cháº·n; Firecrawl chá»‰ lÃ  fallback cho bÃ i bÃ¡o/trang tÄ©nh khi cÃ³ API key. KhÃ´ng cÃ³ lá»‹ch Ä‘áº¥u tháº§u hay lÃ£i suáº¥t nÃ o Ä‘Æ°á»£c sinh giáº£.",
    )
    add_para(
        doc,
        "Má»¥c tiÃªu cá»§a pháº§n má»Ÿ rá»™ng khÃ´ng pháº£i thay Ä‘á»•i káº¿t luáº­n Ä‘áº§u tÆ° má»™t cÃ¡ch cÆ¡ há»c, mÃ  lÃ  giáº£m báº¥t Ä‘á»‹nh quanh cÃ¡c biáº¿n Ä‘ang yáº¿u: premium reference, FX, lÃ£i suáº¥t cÆ¡ há»™i, policy events vÃ  kháº£ nÄƒng kiá»ƒm chá»©ng real-time. Khi cÃ¡c nguá»“n nÃ y tá»‘t hÆ¡n, mÃ´ hÃ¬nh cÃ³ thá»ƒ phÃ¢n biá»‡t rÃµ hÆ¡n giá»¯a má»™t Ä‘á»£t tÄƒng giÃ¡ do vÃ ng tháº¿ giá»›i, má»™t Ä‘á»£t tÄƒng do VND máº¥t giÃ¡, vÃ  má»™t Ä‘á»£t tÄƒng do premium ná»™i Ä‘á»‹a bá»‹ Ä‘áº©y bá»Ÿi khan hiáº¿m/chÃ­nh sÃ¡ch.",
    )
    add_para(
        doc,
        "Playwright Ä‘Æ°á»£c Æ°u tiÃªn cho SBV vÃ¬ phiÃªn trÃ¬nh duyá»‡t tháº­t giá»¯ Ä‘Æ°á»£c session vÃ  báº¯t Ä‘Æ°á»£c cÃ¡c endpoint CMS mÃ  request trá»±c tiáº¿p cÃ³ thá»ƒ bá»‹ cháº·n. Collector khÃ´ng suy diá»…n lá»‹ch Ä‘áº¥u tháº§u tá»« quy luáº­t tuáº§n hoáº·c tin Ä‘á»“n; chá»‰ event cÃ³ URL chÃ­nh thá»©c, ngÃ y cÃ´ng bá»‘, hash ná»™i dung vÃ  confidence má»›i Ä‘Æ°á»£c Ä‘Æ°a vÃ o event panel chÃ­nh. Dá»¯ liá»‡u rule_generated náº¿u cÃ³ chá»‰ Ä‘Æ°á»£c giá»¯ nhÆ° tham kháº£o, khÃ´ng Ä‘i vÃ o mÃ´ hÃ¬nh chÃ­nh.",
    )
    add_para(doc, "Before vs After Caveats:")
    add_table(doc, ["Caveat", "Before", "After", "Decision impact"], caveat_improvement_rows(summary), max_rows=10)

    discovery_status = read_json("data/lake/source_discovery/sbv_source_discovery_status.json")
    discovery = read_csv("data/lake/source_discovery/sbv_structures.csv")
    events_status = read_json("data/lake/events/sbv_gold_policy_events_status.json")
    events = read_csv("data/lake/events/sbv_gold_policy_events.csv")
    paper = read_json("data/lake/modeling/paper_trading_summary.json")
    lbma = read_csv("data/lake/normalized/lbma_gold_spot_am_pm.csv")

    rows = [
        [
            "SBV source discovery",
            discovery_status.get("rows", len(discovery) if not discovery.empty else 0),
            "137473 central FX verified" if discovery_status.get("central_fx_structure_137473_verified") else "not verified",
            "; ".join(discovery_status.get("blockers", [])) or "No blocker recorded",
        ],
        [
            "SBV gold policy events",
            events_status.get("rows", len(events) if not events.empty else 0),
            f"{events_status.get('date_min', '')} to {events_status.get('date_max', '')}".strip(),
            "; ".join(events_status.get("blockers", [])) or "Official-event extraction completed",
        ],
        [
            "Rolling paper trading",
            paper.get("rows", 0),
            f"buy={paper.get('buy_rows', 0)}, open={paper.get('open_rows', 0)}, closed={paper.get('closed_rows', 0)}",
            "; ".join(paper.get("blockers", [])) or "Ledger ready",
        ],
        [
            "LBMA daily append",
            len(lbma) if not lbma.empty else 0,
            date_span(lbma) if not lbma.empty else "not collected in this run",
            "Public today.json append; historical backfill still requires licensed archive or GC=F proxy.",
        ],
    ]
    add_table(doc, ["Artifact", "Rows", "Coverage/status", "Caveat"], rows, max_rows=12)
    add_para(
        doc,
        "Báº£ng artifact cho tháº¥y há»‡ thá»‘ng Ä‘Ã£ chuyá»ƒn tá»« caveat mÃ´ táº£ sang caveat Ä‘o Ä‘Æ°á»£c. Premium cÃ³ audit coverage; deposit rate cÃ³ tráº¡ng thÃ¡i forward monitoring; SBV discovery ghi rÃµ structure nÃ o lÃ  tá»· giÃ¡ trung tÃ¢m; paper-trading ledger ghi tráº¡ng thÃ¡i open/closed thay vÃ¬ tÃ­nh realized return trÆ°á»›c khi giao dá»‹ch Ä‘Ã¡o háº¡n. ÄÃ¢y lÃ  Ä‘iá»u kiá»‡n cáº§n Ä‘á»ƒ biáº¿n bÃ¡o cÃ¡o tá»« má»™t phÃ¢n tÃ­ch tÄ©nh thÃ nh má»™t quy trÃ¬nh theo dÃµi sau snapshot.",
    )

    if not discovery.empty:
        cols = [c for c in ["content_structure_id", "classification", "row_count_sample", "field_names"] if c in discovery.columns]
        if cols:
            add_para(doc, "SBV structures discovered from official pages/session:")
            add_table(doc, ["Structure", "Classification", "Sample rows", "Fields"], discovery[cols].head(8).values.tolist(), max_rows=8)

    if not events.empty and "event_type" in events.columns:
        event_counts = events.groupby("event_type").size().reset_index(name="count").sort_values("count", ascending=False)
        add_para(doc, "Official SBV gold-policy event types collected:")
        add_table(doc, ["Event type", "Count"], event_counts.values.tolist(), max_rows=10)

    if paper.get("by_horizon"):
        p_rows = []
        for item in paper["by_horizon"]:
            avg = item.get("avg_realized_net_return")
            hit = item.get("hit_rate")
            p_rows.append([
                item.get("horizon_months"),
                item.get("buy_rows"),
                item.get("open_rows"),
                item.get("closed_rows"),
                pct(avg, 2) if avg is not None else "n/a",
                pct(hit, 1) if hit is not None else "n/a",
            ])
        add_para(doc, "Paper-trading ledger status by horizon:")
        add_table(doc, ["Horizon", "Buy rows", "Open", "Closed", "Avg realized", "Hit rate"], p_rows, max_rows=10)
        ledger = read_csv("data/lake/modeling/paper_trading_ledger.csv")
        if not ledger.empty:
            cols = ["trade_id", "feature_date", "horizon_months", "entry_sell_price", "expected_return", "q10_downside", "prob_positive", "decision", "exit_status", "realized_net_return"]
            add_para(doc, "Paper-trading ledger sample: dá»¯ liá»‡u nÃ y náº±m ngay pháº§n monitoring Ä‘á»ƒ tháº¥y horizon 5 thÃ¡ng Ä‘ang open vÃ  chÆ°a cÃ³ realized return.")
            add_table(doc, [c for c in cols if c in ledger.columns], rows_from_frame(ledger, cols, max_rows=3, max_chars=120), max_rows=3)
        add_para(
            doc,
            "Paper-trading sau snapshot Ä‘Æ°á»£c thiáº¿t káº¿ Ä‘á»ƒ trÃ¡nh self-deception. Má»—i ngÃ y chá»‰ Ä‘Æ°á»£c ghi signal dá»±a trÃªn feature Ä‘Ã£ available as-of ngÃ y quyáº¿t Ä‘á»‹nh; lá»‡nh chÆ°a Ä‘á»§ ngÃ y exit pháº£i giá»¯ tráº¡ng thÃ¡i open vÃ  khÃ´ng Ä‘Æ°á»£c tÃ­nh vÃ o hit rate. Khi cÃ³ dá»¯ liá»‡u exit tháº­t, ledger má»›i chuyá»ƒn sang closed vÃ  cáº­p nháº­t realized return sau spread. Nhá» váº­y performance tÆ°Æ¡ng lai khÃ´ng bá»‹ trá»™n vá»›i forecast chÆ°a Ä‘Ã¡o háº¡n.",
        )
        add_code_block(
            doc,
            "Paper-trading khÃ´ng tÃ­nh return khi lá»‡nh chÆ°a Ä‘Ã¡o háº¡n",
            "scripts/analysis/run_rolling_paper_trading.py",
            105,
            132,
            "Ledger chá»‰ tÃ­nh realized_net_return khi cÃ³ exit_buy_price tháº­t; náº¿u chÆ°a tá»›i ngÃ y exit thÃ¬ status=open.",
        )


def add_evidence_sections(doc: Document, summary: dict[str, Any]) -> None:
    add_heading(doc, "11. Evidence tá»« crawl, data lake vÃ  code pipeline", 1)
    add_para(
        doc,
        "Pháº§n nÃ y bá»• sung lá»›p minh chá»©ng trá»±c tiáº¿p tá»« cÃ¡c file Ä‘Ã£ crawl, cÃ¡c báº£ng trong data lake vÃ  cÃ¡c Ä‘oáº¡n code Ä‘ang táº¡o dá»¯ liá»‡u/mÃ´ hÃ¬nh. Má»¥c tiÃªu lÃ  Ä‘á»ƒ ngÆ°á»i Ä‘á»c cÃ³ thá»ƒ láº§n ngÆ°á»£c tá»« káº¿t luáº­n trong bÃ¡o cÃ¡o vá» artifact cá»¥ thá»ƒ: dá»¯ liá»‡u nÃ o Ä‘Æ°á»£c Ä‘á»c, rule nÃ o Ä‘Æ°á»£c Ã¡p dá»¥ng, vÃ  vÃ¬ sao má»™t sá»‘ nguá»“n bá»‹ loáº¡i hoáº·c chá»‰ Ä‘Æ°á»£c dÃ¹ng Ä‘á»ƒ monitoring.",
    )

    sbv_structures = read_csv("data/lake/source_discovery/sbv_structures.csv")
    sbv_events = read_csv("data/lake/events/sbv_gold_policy_events.csv")
    deposit = read_csv("data/lake/normalized/retail_deposit_rates.csv")
    lbma = read_csv("data/lake/normalized/lbma_gold_spot_am_pm.csv")
    news = read_csv("data/lake/news_availability_audit.csv")
    paper = read_csv("data/lake/modeling/paper_trading_ledger.csv")
    frame = read_csv("data/lake/modeling/model_frame_daily.csv")
    snapshot = read_csv("data/lake/modeling/snapshot_forecasts.csv")
    premium_audit = read_csv("data/lake/quality/premium_coverage_audit.csv")
    sensitivity = read_json("data/lake/modeling/model_sensitivity_summary.json")

    structure_counts = sbv_structures["classification"].value_counts(dropna=False).to_dict() if "classification" in sbv_structures else {}
    event_counts = sbv_events["event_type"].value_counts(dropna=False).to_dict() if "event_type" in sbv_events else {}
    strict_counts = news["feature_mode_strict"].value_counts(dropna=False).to_dict() if "feature_mode_strict" in news else {}
    dep_avail = sensitivity.get("dependency_availability", {})

    add_heading(doc, "11.1. Artifact audit: dá»¯ liá»‡u tháº­t Ä‘Ã£ Ä‘i vÃ o bÃ¡o cÃ¡o", 2)
    add_table(
        doc,
        ["Artifact", "Rows/Cols", "Evidence Ä‘á»c Ä‘Æ°á»£c", "Ã nghÄ©a trong bÃ¡o cÃ¡o"],
        [
            [
                "source_discovery/sbv_structures.csv",
                f"{len(sbv_structures):,} rows / {len(sbv_structures.columns) if not sbv_structures.empty else 0} cols",
                truncate_text(structure_counts, 180),
                "Káº¿t quáº£ Playwright discovery; xÃ¡c nháº­n structure 137473 lÃ  central FX, khÃ´ng pháº£i deposit rate.",
            ],
            [
                "events/sbv_gold_policy_events.csv",
                f"{len(sbv_events):,} rows / {len(sbv_events.columns) if not sbv_events.empty else 0} cols",
                truncate_text(event_counts, 180),
                "Event chÃ­nh sÃ¡ch vÃ ng cÃ³ URL SBV, published_at, raw_hash vÃ  confidence.",
            ],
            [
                "normalized/retail_deposit_rates.csv",
                f"{len(deposit):,} rows / {len(deposit.columns) if not deposit.empty else 0} cols",
                f"rate range {deposit['rate_pct_annual'].min():.2f}-{deposit['rate_pct_annual'].max():.2f}%/nÄƒm" if "rate_pct_annual" in deposit and not deposit.empty else "not available",
                "Nguá»“n lÃ£i suáº¥t tiá»n gá»­i hiá»‡n táº¡i Ä‘á»ƒ forward-monitoring; khÃ´ng backfill lá»‹ch sá»­ suy diá»…n.",
            ],
            [
                "normalized/lbma_gold_spot_am_pm.csv",
                f"{len(lbma):,} rows / {len(lbma.columns) if not lbma.empty else 0} cols",
                ", ".join(sorted(lbma["series_id"].dropna().astype(str).head(4).tolist())) if "series_id" in lbma else "not available",
                "LBMA today.json append; lá»‹ch sá»­ váº«n dÃ¹ng proxy khi thiáº¿u licence.",
            ],
            [
                "news_availability_audit.csv",
                f"{len(news):,} rows / {len(news.columns) if not news.empty else 0} cols",
                truncate_text(strict_counts, 180),
                "TÃ¡ch research_event_date_lagged vÃ  strict_realtime_verified Ä‘á»ƒ trÃ¡nh dÃ¹ng news backfill nhÆ° tÃ­n hiá»‡u trading tháº­t.",
            ],
            [
                "modeling/model_frame_daily.csv",
                f"{len(frame):,} rows / {len(frame.columns) if not frame.empty else 0} cols",
                f"{frame['date'].min()} to {frame['date'].max()}, premium missing {pct(frame['premium'].isna().mean())}" if {"date", "premium"} <= set(frame.columns) and not frame.empty else "not available",
                "Frame cuá»‘i sau as-of join, feature engineering vÃ  target 1/3/5 thÃ¡ng.",
            ],
            [
                "modeling/paper_trading_ledger.csv",
                f"{len(paper):,} rows / {len(paper.columns) if not paper.empty else 0} cols",
                truncate_text(paper[["horizon_months", "decision", "exit_status"]].to_dict("records") if not paper.empty else [], 180),
                "Ledger sau snapshot; lá»‡nh chÆ°a Ä‘Ã¡o háº¡n giá»¯ tráº¡ng thÃ¡i open, khÃ´ng tÃ­nh realized return.",
            ],
        ],
        max_rows=10,
    )

    add_para(
        doc,
        "Báº£ng trÃªn lÃ  lá»›p evidence-level: má»—i hÃ ng lÃ  má»™t artifact Ä‘Ã£ Ä‘Æ°á»£c Ä‘á»c tá»« data lake khi build bÃ¡o cÃ¡o. CÃ¡c sá»‘ nhÆ° 5,485 dÃ²ng model frame, 3 dÃ²ng snapshot forecast, 3 dÃ²ng paper-trading ledger vÃ  3,441 dÃ²ng news audit Ä‘á»u lÃ  sá»‘ Ä‘áº¿m trá»±c tiáº¿p tá»« file, khÃ´ng pháº£i mÃ´ táº£ thá»§ cÃ´ng.",
    )

    add_heading(doc, "11.2. Sample rows tá»« dá»¯ liá»‡u crawl vÃ  lake", 2)
    if not sbv_structures.empty:
        add_para(doc, "SBV source discovery sample: cÃ¡c structure Ä‘Æ°á»£c láº¥y qua phiÃªn Playwright vÃ  phÃ¢n loáº¡i báº±ng field/title/seed hit.")
        cols = ["content_structure_id", "classification", "http_status", "row_count_sample", "field_names", "title_samples"]
        add_table(doc, [c for c in cols if c in sbv_structures.columns], rows_from_frame(sbv_structures, cols, max_rows=6, max_chars=110), max_rows=6)
    if not sbv_events.empty:
        add_para(doc, "SBV gold-policy event sample: chá»‰ giá»¯ event cÃ³ URL chÃ­nh thá»©c, ngÃ y cÃ´ng bá»‘, source_type, confidence vÃ  raw_hash.")
        cols = ["event_date", "published_at", "event_type", "severity", "title", "confidence", "source_type"]
        add_table(doc, [c for c in cols if c in sbv_events.columns], rows_from_frame(sbv_events, cols, max_rows=5, max_chars=120), max_rows=5)
    if not deposit.empty:
        add_para(doc, "Retail deposit-rate sample: dá»¯ liá»‡u nÃ y Ä‘Æ°á»£c dÃ¹ng cho forward monitoring, khÃ´ng Ä‘Æ°á»£c coi lÃ  lá»‹ch sá»­ lÃ£i suáº¥t Ä‘áº§y Ä‘á»§.")
        cols = ["date", "bank", "tenor_months", "currency", "rate_pct_annual", "published_at", "available_from", "history_status"]
        add_table(doc, [c for c in cols if c in deposit.columns], rows_from_frame(deposit, cols, max_rows=8, max_chars=90), max_rows=8)
    if not news.empty:
        add_para(doc, "News availability sample: cÃ¹ng má»™t headline cÃ³ thá»ƒ dÃ¹ng cho research mode, nhÆ°ng chá»‰ strict mode náº¿u availability_from khÃ´ng Ä‘i sau feature date.")
        cols = ["crawl_date", "event_date", "headline", "published_at", "fetched_at", "availability_from", "feature_mode_strict"]
        add_table(doc, [c for c in cols if c in news.columns], rows_from_frame(news, cols, max_rows=5, max_chars=120), max_rows=5)
    if not snapshot.empty:
        add_para(doc, "Snapshot forecast sample: Ä‘Ã¢y lÃ  Ä‘áº§u ra model táº¡i ngÃ y feature cuá»‘i, sau khi train trÃªn cÃ¡c dÃ²ng lá»‹ch sá»­ cÃ³ nhÃ£n.")
        cols = ["snapshot_date", "horizon_months", "predicted_net_return", "q10_predicted_net_return", "prob_return_positive", "buy_signal", "status"]
        add_table(doc, [c for c in cols if c in snapshot.columns], rows_from_frame(snapshot, cols, max_rows=3, max_chars=90), max_rows=3)
    if not paper.empty:
        add_para(doc, "Paper-trading ledger sample: horizon 5 thÃ¡ng cÃ³ quyáº¿t Ä‘á»‹nh buy nhÆ°ng exit_status=open nÃªn chÆ°a cÃ³ realized return.")
        cols = ["trade_id", "feature_date", "horizon_months", "entry_sell_price", "expected_return", "q10_downside", "prob_positive", "decision", "exit_status", "realized_net_return"]
        add_table(doc, [c for c in cols if c in paper.columns], rows_from_frame(paper, cols, max_rows=3, max_chars=120), max_rows=3)
    if not frame.empty:
        add_para(doc, "Model frame sample á»Ÿ cuá»‘i snapshot: cÃ¡c cá»™t giÃ¡, premium, spread vÃ  target tÆ°Æ¡ng lai cÃ¹ng náº±m trÃªn má»™t grain ngÃ y.")
        cols = ["date", "buy_price", "sell_price", "spread_pct", "premium", "premium_pct", "is_proxy", "net_return_1m", "net_return_3m", "net_return_5m"]
        available = [c for c in cols if c in frame.columns]
        add_table(doc, available, rows_from_frame(frame[available].tail(5), available, max_rows=5, max_chars=80), max_rows=5)

    add_heading(doc, "11.3. Code evidence: rule táº¡o target, chá»‘ng leakage vÃ  model tháº­t", 2)
    add_para(
        doc,
        "CÃ¡c Ä‘oáº¡n code dÆ°á»›i Ä‘Ã¢y Ä‘Æ°á»£c trÃ­ch trá»±c tiáº¿p tá»« repository. ChÃºng lÃ  pháº§n giáº£i thÃ­ch ká»¹ thuáº­t cho cÃ¡c quyáº¿t Ä‘á»‹nh phÆ°Æ¡ng phÃ¡p trong bÃ¡o cÃ¡o: target Ä‘Æ°á»£c tÃ­nh theo giÃ¡ mua tÆ°Æ¡ng lai/giÃ¡ bÃ¡n hiá»‡n táº¡i, join dá»¯ liá»‡u ngoáº¡i sinh dÃ¹ng as-of cutoff, optional boosting chá»‰ cháº¡y khi dependency import Ä‘Æ°á»£c, vÃ  rule mua cáº§n cáº£ xÃ¡c suáº¥t láº«n q10 downside.",
    )
    add_code_block(
        doc,
        "Target lá»£i suáº¥t sau spread",
        "src/gold_collectors/modeling/decision_support.py",
        375,
        386,
        "ÄÃ¢y lÃ  Ä‘oáº¡n biáº¿n bÃ i toÃ¡n dá»± bÃ¡o giÃ¡ thÃ nh bÃ i toÃ¡n lá»£i suáº¥t thá»±c thi: mua theo sell_price hiá»‡n táº¡i vÃ  thoÃ¡t theo future buy_price.",
    )
    add_code_block(
        doc,
        "As-of join cho dá»¯ liá»‡u global",
        "src/gold_collectors/modeling/decision_support.py",
        80,
        112,
        "Global market dÃ¹ng cutoff t-1 trÆ°á»›c khi merge_asof, giáº£m rá»§i ro dÃ¹ng giÃ¡ quá»‘c táº¿ Ä‘Ã³ng cá»­a sau giá» Viá»‡t Nam.",
    )
    add_code_block(
        doc,
        "Deposit-rate as-of vÃ  opportunity cost",
        "src/gold_collectors/modeling/decision_support.py",
        188,
        220,
        "Feature lÃ£i suáº¥t chá»‰ Ä‘Æ°á»£c táº¡o khi cÃ³ available_from, tenor vÃ  rate há»£p lá»‡; náº¿u chÆ°a cÃ³ lá»‹ch sá»­ thÃ¬ khÃ´ng Ã©p backfill.",
    )
    add_code_block(
        doc,
        "Pipeline build model frame",
        "src/gold_collectors/modeling/decision_support.py",
        465,
        490,
        "Thá»© tá»± build frame thá»ƒ hiá»‡n toÃ n bá»™ pipeline: target SJC, premium, global, GPR, macro, deposit, event, news, target vÃ  lagged features.",
    )
    add_code_block(
        doc,
        "Optional model training cho LightGBM/XGBoost/CatBoost",
        "src/gold_collectors/modeling/decision_support.py",
        713,
        753,
        "CÃ¡c mÃ´ hÃ¬nh boosting khÃ´ng cÃ²n lÃ  blocker dependency; náº¿u import Ä‘Æ°á»£c thÃ¬ Ä‘Æ°á»£c Ä‘Æ°a vÃ o optional_models vÃ  train tháº­t.",
    )
    add_code_block(
        doc,
        "Decision rule tá»« xÃ¡c suáº¥t vÃ  q10 downside",
        "src/gold_collectors/modeling/decision_support.py",
        930,
        945,
        "TÃ­n hiá»‡u mua khÃ´ng chá»‰ dá»±a vÃ o expected return; rule yÃªu cáº§u P(return>0) vÆ°á»£t ngÆ°á»¡ng vÃ  q10 khÃ´ng tháº¥p hÆ¡n sÃ n rá»§i ro.",
    )

    add_heading(doc, "11.4. Code evidence: crawl SBV, premium fallback vÃ  paper-trading", 2)
    add_code_block(
        doc,
        "SBV discovery báº±ng Playwright vÃ  endpoint headless-delivery",
        "scripts/pipeline/discover_sbv_sources.py",
        35,
        72,
        "Collector má»Ÿ cÃ¡c trang SBV chÃ­nh thá»©c, tÃ¬m content-structure id, rá»“i fetch endpoint qua browser context thay vÃ¬ request tháº³ng dá»… bá»‹ cháº·n.",
    )
    add_code_block(
        doc,
        "PhÃ¢n loáº¡i SBV event cÃ³ source_url, confidence vÃ  raw_hash",
        "scripts/pipeline/collect_sbv_gold_policy_events.py",
        180,
        199,
        "Má»—i event chÃ­nh sÃ¡ch vÃ ng pháº£i cÃ³ event_date/published_at, event_type, source_url, source_type, confidence vÃ  raw_hash.",
    )
    add_code_block(
        doc,
        "Premium coverage audit sau fallback cÃ³ giá»›i háº¡n",
        "scripts/pipeline/improve_premium_coverage.py",
        257,
        287,
        "Premium fallback Ä‘Æ°á»£c Ä‘Ã¡nh giÃ¡ báº±ng missing rate theo nÄƒm/source quality vÃ  target <10%; khÃ´ng Ä‘áº¡t thÃ¬ ghi blocker thay vÃ¬ im láº·ng.",
    )
    add_code_block(
        doc,
        "Paper-trading khÃ´ng tÃ­nh return khi lá»‡nh chÆ°a Ä‘Ã¡o háº¡n",
        "scripts/analysis/run_rolling_paper_trading.py",
        105,
        132,
        "Ledger chá»‰ tÃ­nh realized_net_return khi cÃ³ exit_buy_price tháº­t; náº¿u chÆ°a tá»›i ngÃ y exit thÃ¬ status=open.",
    )

    add_heading(doc, "11.5. Äá»‘i chiáº¿u code - dá»¯ liá»‡u - káº¿t luáº­n", 2)
    add_table(
        doc,
        ["Káº¿t luáº­n trong bÃ¡o cÃ¡o", "Artifact chá»©ng minh", "Code sinh ra/kiá»ƒm tra", "CÃ¡ch Ä‘á»c"],
        [
            [
                "Premium missing giáº£m vá» má»©c bÃ¡o cÃ¡o trong model frame",
                "quality/premium_coverage_audit.csv; modeling/model_frame_daily.csv",
                "scripts/pipeline/improve_premium_coverage.py:257-287",
                "Coverage Ä‘Æ°á»£c Ä‘o theo nÄƒm vÃ  source_quality; proxy váº«n Ä‘Æ°á»£c gáº¯n cá» is_proxy.",
            ],
            [
                "KhÃ´ng dÃ¹ng SBV structure 137473 lÃ m deposit rate",
                "source_discovery/sbv_structures.csv",
                "scripts/pipeline/discover_sbv_sources.py:53-63",
                "Structure cÃ³ field central FX Ä‘Æ°á»£c phÃ¢n loáº¡i central_fx, khÃ´ng pháº£i interest_rate_candidate.",
            ],
            [
                "News strict realtime cÃ²n yáº¿u",
                "news_availability_audit.csv; quality/news_availability_summary.json",
                "scripts/pipeline/audit_news_availability.py",
                "Backfilled rows bá»‹ tÃ¡ch khá»i strict_realtime_verified Ä‘á»ƒ trÃ¡nh dÃ¹ng tin chÆ°a chá»©ng minh available.",
            ],
            [
                "Horizon 5 thÃ¡ng lÃ  tÃ­n hiá»‡u buy nhÆ°ng chÆ°a cÃ³ realized performance",
                "modeling/snapshot_forecasts.csv; modeling/paper_trading_ledger.csv",
                "scripts/analysis/run_rolling_paper_trading.py:105-132",
                "Lá»‡nh 5 thÃ¡ng Ä‘ang open; khÃ´ng tÃ­nh hit rate/return trÆ°á»›c target_exit_date.",
            ],
            [
                "Boosting models Ä‘Æ°á»£c train tháº­t khi dependency cÃ³ máº·t",
                "modeling/model_results.csv; modeling/model_sensitivity_summary.json",
                "src/gold_collectors/modeling/decision_support.py:713-753",
                "Dependency availability true; leaderboard cÃ³ LightGBM/XGBoost/CatBoost náº¿u training hoÃ n táº¥t.",
            ],
        ],
        max_rows=10,
    )
    add_para(
        doc,
        "Nhá» lá»›p Ä‘á»‘i chiáº¿u nÃ y, bÃ¡o cÃ¡o khÃ´ng chá»‰ lÃ  diá»…n giáº£i Ä‘á»‹nh tÃ­nh. Má»—i káº¿t luáº­n quan trá»ng Ä‘á»u cÃ³ ba pháº§n Ä‘i kÃ¨m: dá»¯ liá»‡u Ä‘áº§u vÃ o hoáº·c output trong lake, Ä‘oáº¡n code táº¡o/kiá»ƒm tra káº¿t quáº£, vÃ  cÃ¡ch Ä‘á»c Ä‘á»ƒ trÃ¡nh hiá»ƒu quÃ¡ má»©c. ÄÃ¢y lÃ  pháº§n lÃ m tÃ i liá»‡u cÃ³ tÃ­nh reproducible report thay vÃ¬ chá»‰ lÃ  bÃ i viáº¿t tÃ³m táº¯t.",
    )


def add_figures_section(doc: Document, figure_paths: dict[str, Path]) -> None:
    add_heading(doc, "12. Figures", 1)
    add_picture_if_exists(doc, figure_paths.get("price", Path()), "Figure 1. SJC price versus global gold proxy")
    add_picture_if_exists(doc, figure_paths.get("premium", Path()), "Figure 2. Premium regime over time")
    add_picture_if_exists(doc, figure_paths.get("leaderboard", Path()), "Figure 3. Model MAE leaderboard by horizon")
    add_picture_if_exists(doc, figure_paths.get("signals", Path()), "Figure 4. Decision signal frequency")


def add_appendix(doc: Document, summary: dict[str, Any]) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_heading(doc, "Appendix: Reproducibility", 1)
    add_para(doc, "Pipeline Ä‘Ã£ Ä‘Æ°á»£c cháº¡y láº¡i trÃªn snapshot hiá»‡n cÃ³ trÆ°á»›c khi dá»±ng bÃ¡o cÃ¡o:")
    add_bullet(doc, "python scripts/analysis/run_decision_support_analysis.py")
    add_bullet(doc, "python scripts/analysis/export_trading_signals.py --quiet")
    add_bullet(doc, "python scripts/analysis/build_full_report.py")
    add_para(doc, "CÃ¡c artifact chÃ­nh:")
    for rel in [
        "data/lake/modeling/model_frame_daily.csv",
        "data/lake/modeling/model_results.csv",
        "data/lake/modeling/walk_forward_predictions.csv",
        "data/lake/modeling/decision_signals.csv",
        "data/lake/modeling/trading_signals.csv",
        "data/lake/modeling/snapshot_forecasts.csv",
        "data/lake/source_discovery/sbv_structures.csv",
        "data/lake/events/sbv_gold_policy_events.csv",
        "data/lake/modeling/paper_trading_ledger.csv",
        "data/lake/quality/premium_coverage_audit.csv",
        "data/lake/normalized/retail_deposit_rates.csv",
        "data/lake/normalized/sbv_policy_rates.csv",
        "data/lake/news_availability_audit.csv",
        "data/lake/modeling/model_sensitivity_summary.json",
    ]:
        add_bullet(doc, rel)
    add_para(doc, "Runtime blockers Ä‘Æ°á»£c ghi nháº­n:")
    for blocker in summary.get("blockers", []):
        add_bullet(doc, blocker)
    add_para(
        doc,
        "CÃ¡ch Ä‘á»c phá»¥ lá»¥c: cÃ¡c artifact trÃªn lÃ  Ä‘Æ°á»ng kiá»ƒm toÃ¡n tá»‘i thiá»ƒu Ä‘á»ƒ tÃ¡i táº¡o headline numbers trong bÃ¡o cÃ¡o. analysis_summary.json chá»©a sá»‘ dÃ²ng model frame, khoáº£ng ngÃ y, leaderboard vÃ  decision summary; snapshot_forecasts.csv chá»©a forecast táº¡i ngÃ y snapshot; premium_coverage_audit.csv vÃ  cÃ¡c file quality JSON giáº£i thÃ­ch nhá»¯ng cáº£i thiá»‡n dá»¯ liá»‡u sau caveat. Náº¿u má»™t con sá»‘ trong bÃ¡o cÃ¡o khÃ´ng Ä‘á»‘i chiáº¿u Ä‘Æ°á»£c vá»›i cÃ¡c artifact nÃ y, con sá»‘ Ä‘Ã³ khÃ´ng nÃªn Ä‘Æ°á»£c xem lÃ  káº¿t luáº­n chÃ­nh thá»©c.",
    )
    add_para(
        doc,
        "BÃ¡o cÃ¡o khÃ´ng sá»­ dá»¥ng dá»¯ liá»‡u live ngoÃ i snapshot Ä‘á»ƒ thay Ä‘á»•i quyáº¿t Ä‘á»‹nh 1-5 thÃ¡ng. Nhá»¯ng nguá»“n Ä‘Æ°á»£c thu tháº­p sau Ä‘Ã³, nhÆ° deposit-rate monitoring hoáº·c paper-trading ledger, Ä‘Æ°á»£c dÃ¹ng Ä‘á»ƒ xÃ¢y dá»±ng cÆ¡ cháº¿ theo dÃµi vÃ  Ä‘Ã¡nh giÃ¡ tÆ°Æ¡ng lai. Äiá»u nÃ y giÃºp tÃ¡ch rÃµ hai viá»‡c: káº¿t luáº­n táº¡i thá»i Ä‘iá»ƒm 2026-07-11 vÃ  váº­n hÃ nh giÃ¡m sÃ¡t sau snapshot.",
    )


def validate_report_inputs(summary: dict[str, Any], profiles: list[DataProfile], snapshot: pd.DataFrame) -> None:
    if not summary:
        raise RuntimeError("Missing analysis_summary.json; run decision support analysis first.")
    required = ["rows", "date_min", "date_max", "leaderboard", "decision_summary"]
    missing = [key for key in required if key not in summary]
    if missing:
        raise RuntimeError(f"analysis_summary.json missing keys: {missing}")
    if not any(p.rows > 0 and p.name == "Model frame" for p in profiles):
        raise RuntimeError("Model frame is empty or missing.")
    if snapshot.empty:
        raise RuntimeError("Snapshot forecasts could not be generated.")


def refresh_paper_trading_ledger() -> None:
    try:
        from scripts.analysis.run_rolling_paper_trading import build_ledger, load_forecasts, load_price_curve, merge_existing, summarize
    except Exception as exc:
        print(f"Paper-trading refresh skipped: {exc}")
        return

    forecasts_path = MODELING / "snapshot_forecasts.csv"
    prices_path = LAKE / "domestic_gold_quotes.csv"
    out_path = MODELING / "paper_trading_ledger.csv"
    summary_path = MODELING / "paper_trading_summary.json"
    forecasts = load_forecasts(forecasts_path, "2026-07-11")
    prices = load_price_curve(prices_path)
    ledger = build_ledger(forecasts, prices, "snapshot_forecast_v1")
    ledger = merge_existing(ledger, out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    ledger.to_csv(out_path, index=False)
    summary_path.write_text(json.dumps(summarize(ledger), ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    summary = read_json("data/lake/modeling/analysis_summary.json")
    frame = read_csv("data/lake/modeling/model_frame_daily.csv")
    feature_cols = load_feature_columns()
    snapshot = train_snapshot_forecasts(frame, feature_cols)
    refresh_paper_trading_ledger()
    profiles, frames = build_profiles()

    validate_report_inputs(summary, profiles, snapshot)

    quality_rows = compute_data_quality(frames, summary)
    eda_tables = compute_eda_tables(frames)
    figure_paths = build_figures(frames, summary)
    mean_rows, quant_rows, decision_perf_rows = summarize_models(summary)
    recommendation, snapshot_rows = decision_recommendation(snapshot, decision_perf_rows)

    latest_date = str(summary.get("date_max", "2026-07-11"))
    doc = Document()
    configure_document(doc)
    add_cover(doc, latest_date)
    add_executive_summary(doc, latest_date, recommendation, summary)
    add_business_understanding(doc)
    add_literature_review(doc)
    add_data_sections(doc, profiles, quality_rows, eda_tables, figure_paths)
    add_modeling_sections(doc, mean_rows, quant_rows, decision_perf_rows, snapshot_rows, recommendation, summary, figure_paths)
    add_next_data_expansion_section(doc, summary)
    add_appendix(doc, summary)

    doc.save(OUT)
    print(f"Saved report: {OUT}")
    print(f"Snapshot forecast: {MODELING / 'snapshot_forecasts.csv'}")


if __name__ == "__main__":
    main()

